from io import BytesIO
from datetime import datetime
import html
import os
import re

import pandas as pd
import streamlit as st

from docx import Document
from docx.shared import RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ============================================================
# PAGE CONFIG
# ============================================================

st.set_page_config(
    page_title="Strategic District Field Guide",
    page_icon="📘",
    layout="wide",
    initial_sidebar_state="expanded",
)

BUILT_IN_WORKBOOK = "Texas Top 24 Research.xlsx"

SHEET_STRATEGIC = "Strategic Indicators"
SHEET_SCORECARD = "Master Scorecard"
SHEET_BASIC = "Basic District Info"
SHEET_CONTACTS = "District Contacts"
SHEET_LEADERSHIP = "Leadership and Governance"
SHEET_CSI = "CSI"
SHEET_TSI = "TSI"
DISTRICT_COL = "District Name"


# ============================================================
# STYLING
# ============================================================

CUSTOM_CSS = """
<style>
:root {
    --navy:#102a43;
    --blue:#1d4ed8;
    --blue-soft:#eff6ff;
    --blue-border:#bfdbfe;
    --green:#166534;
    --green-soft:#f0fdf4;
    --red:#991b1b;
    --red-soft:#fef2f2;
    --amber:#92400e;
    --amber-soft:#fffbeb;
    --slate:#334155;
    --slate-soft:#f8fafc;
    --page:#f3f6fa;
    --surface:#ffffff;
    --text:#0f172a;
    --muted:#475569;
    --border:#cbd5e1;
}

.stApp,
[data-testid="stAppViewContainer"],
[data-testid="stHeader"] {
    background: var(--page) !important;
    color: var(--text) !important;
}

html, body, p, li, span, div, label, section, article,
.stMarkdown, .stMarkdown p, .stMarkdown li, .stMarkdown span, .stMarkdown div {
    color: var(--text) !important;
}

.block-container {
    padding-top: 1.2rem;
    padding-bottom: 2.5rem;
    max-width: 1220px;
}

[data-testid="stSidebar"] {
    background: #ffffff !important;
    border-right: 1px solid var(--border) !important;
}

[data-testid="stSidebar"] * {
    color: var(--text) !important;
}

h1, h2, h3, h4, h5, h6 {
    color: var(--navy) !important;
}

.main-title {
    font-size: 2rem;
    font-weight: 850;
    color: var(--navy) !important;
    margin-bottom: .25rem;
    letter-spacing:-.02em;
}

.subtitle {
    color: var(--muted) !important;
    font-size: .98rem;
    margin-bottom: 1rem;
}

.section-title {
    color: var(--navy) !important;
    font-size: 1.25rem;
    font-weight: 800;
    margin: .6rem 0 .3rem;
}

.helper-text {
    color: var(--muted) !important;
    font-size: .9rem;
}

.stTabs [data-baseweb="tab-list"] {
    background: #ffffff !important;
    border: 1px solid var(--border) !important;
    border-radius: 14px !important;
    padding: .25rem !important;
}

.stTabs [data-baseweb="tab"] {
    color: var(--slate) !important;
    background: transparent !important;
    border-radius: 10px !important;
    font-weight: 700 !important;
}

.stTabs [aria-selected="true"] {
    background: var(--blue-soft) !important;
    color: var(--blue) !important;
}

[data-testid="stMetric"] {
    background: #ffffff !important;
    border: 1px solid var(--border) !important;
    border-radius: 16px !important;
    padding: .85rem !important;
    box-shadow: 0 4px 14px rgba(15, 23, 42, .05) !important;
}

[data-testid="stMetric"] * {
    color: var(--text) !important;
}

input, textarea, select,
[data-baseweb="input"],
[data-baseweb="select"],
[data-baseweb="textarea"] {
    background: #ffffff !important;
    color: var(--text) !important;
    border-color: var(--border) !important;
}

.stButton button,
.stDownloadButton button {
    background: var(--blue) !important;
    color: #ffffff !important;
    border: 1px solid var(--blue) !important;
    border-radius: 10px !important;
    font-weight: 750 !important;
}

.stButton button:hover,
.stDownloadButton button:hover {
    background: #1e40af !important;
    border-color: #1e40af !important;
    color: #ffffff !important;
}

[data-testid="stExpander"] {
    background: #ffffff !important;
    border: 1px solid var(--border) !important;
    border-radius: 14px !important;
}

.card {
    border:1px solid var(--border);
    border-radius:20px;
    padding:1.05rem;
    margin-bottom:1rem;
    background:var(--surface);
    color: var(--text) !important;
    box-shadow:0 8px 24px rgba(15,23,42,.08);
}

.card, .card * {
    color: var(--text) !important;
}

.card h3 {
    color:var(--navy) !important;
    margin-top:0;
    margin-bottom:.2rem;
    font-size:1.35rem;
    letter-spacing:-.01em;
}

.meta {
    color:var(--muted) !important;
    font-size:.88rem;
    margin-bottom:.75rem;
    line-height:1.45;
}

.summary-grid {
    display:grid;
    grid-template-columns: repeat(3, minmax(0,1fr));
    gap:.75rem;
    margin:.75rem 0 .6rem;
}

.summary-box {
    background:#ffffff;
    border:1px solid var(--border);
    border-radius:14px;
    padding:.75rem;
    box-shadow:0 3px 12px rgba(15,23,42,.04);
}

.summary-label {
    color:var(--muted) !important;
    font-size:.72rem;
    font-weight:850;
    text-transform:uppercase;
    letter-spacing:.05em;
    margin-bottom:.2rem;
}

.summary-value {
    color:var(--text) !important;
    font-size:.92rem;
    line-height:1.35;
}

.quickprep {
    background: #ffffff;
    border:1px solid var(--blue-border);
    border-left:5px solid var(--blue);
    border-radius:16px;
    padding:.9rem;
    margin:.7rem 0 .8rem;
    color: var(--text) !important;
    box-shadow: 0 4px 16px rgba(15, 23, 42, .05);
}

.quickprep-title {
    color:var(--navy) !important;
    font-weight:850;
    font-size:.95rem;
    text-transform:uppercase;
    letter-spacing:.04em;
    margin-bottom:.25rem;
}

.badge, .chip {
    display:inline-block;
    border-radius:999px;
    padding:.25rem .58rem;
    margin:.14rem .16rem .14rem 0;
    font-size:.76rem;
    font-weight:750;
    line-height:1.2;
    border:1px solid transparent;
}

.chip {
    background:var(--slate-soft);
    color:var(--slate) !important;
    border-color:#e2e8f0;
}

.tag-math { background:#dbeafe; color:#1e3a8a !important; border-color:#bfdbfe; }
.tag-mtss { background:#ccfbf1; color:#134e4a !important; border-color:#99f6e4; }
.tag-spedell { background:#ede9fe; color:#4c1d95 !important; border-color:#ddd6fe; }
.tag-ccmr { background:#fef3c7; color:#78350f !important; border-color:#fde68a; }
.tag-teacher { background:#f1f5f9; color:#1e293b !important; border-color:#e2e8f0; }
.tag-hqim { background:#e0e7ff; color:#312e81 !important; border-color:#c7d2fe; }
.tag-funding { background:#fef9c3; color:#713f12 !important; border-color:#fde68a; }
.tag-relationship { background:#dcfce7; color:#14532d !important; border-color:#bbf7d0; }
.tag-literacy { background:#fae8ff; color:#701a75 !important; border-color:#f5d0fe; }
.tag-coaching { background:#ecfccb; color:#365314 !important; border-color:#d9f99d; }
.tag-data { background:#e0f2fe; color:#075985 !important; border-color:#bae6fd; }
.tag-default { background:#eef2ff; color:#312e81 !important; border-color:#e0e7ff; }

.priority {
    display:inline-block;
    border-radius:999px;
    padding:.28rem .62rem;
    font-size:.75rem;
    font-weight:850;
    vertical-align:middle;
}

.priority-very-high { background:var(--red-soft); color:var(--red) !important; border:1px solid #fecaca; }
.priority-high { background:var(--green-soft); color:var(--green) !important; border:1px solid #bbf7d0; }
.priority-medium-high { background:var(--amber-soft); color:var(--amber) !important; border:1px solid #fde68a; }
.priority-medium { background:#e0f2fe; color:#075985 !important; border:1px solid #bae6fd; }

.solution-card {
    background:#ffffff;
    border:1px solid var(--border);
    border-left:5px solid var(--green);
    border-radius:16px;
    padding:.85rem;
    margin:.55rem 0;
    box-shadow:0 4px 16px rgba(15,23,42,.05);
}

.hot-card {
    background:#ffffff;
    border:1px solid var(--border);
    border-left:5px solid var(--red);
    border-radius:16px;
    padding:.85rem;
    margin:.4rem 0;
    box-shadow:0 4px 16px rgba(15,23,42,.05);
}

.howto-box {
    background:#ffffff;
    color: var(--text) !important;
    border:1px solid var(--border);
    border-radius:18px;
    padding:1rem 1.1rem;
    margin:.75rem 0;
    box-shadow:0 4px 16px rgba(15,23,42,.05);
}

@media (max-width: 800px) {
    .summary-grid {
        grid-template-columns: 1fr;
    }
}
</style>
"""

st.markdown(CUSTOM_CSS, unsafe_allow_html=True)


# ============================================================
# BASIC HELPERS
# ============================================================

def safe_html(value):
    return html.escape(str(value))


def normalize_text(value):
    if value is None:
        return ""
    try:
        if pd.isna(value):
            return ""
    except Exception:
        pass
    return str(value).strip()


def normalize_key(value):
    text = normalize_text(value).lower()
    text = text.replace("\n", " ").replace("\r", " ").replace("\xa0", " ")
    text = re.sub(r"\s+", " ", text).strip()
    return text


def district_key(value):
    return normalize_text(value).upper().strip()


def make_unique_columns(columns):
    seen = {}
    output = []
    for col in columns:
        base = normalize_text(col) or "Unnamed"
        if base not in seen:
            seen[base] = 0
            output.append(base)
        else:
            seen[base] += 1
            output.append(f"{base}.{seen[base]}")
    return output


def clean_columns(df):
    if df.empty:
        return df

    df = df.copy()
    df.columns = make_unique_columns(
        [str(c).strip().replace("\n", " ").replace("\r", " ") for c in df.columns]
    )

    drop_cols = [
        c for c in df.columns
        if normalize_key(c).startswith("methods:")
        or normalize_key(c).startswith("unnamed")
    ]

    if drop_cols:
        df = df.drop(columns=drop_cols)

    return df


def find_sheet_name(xls, desired_name):
    lookup = {normalize_key(sheet): sheet for sheet in xls.sheet_names}
    return lookup.get(normalize_key(desired_name))


def read_table_sheet(xls, desired_sheet_name):
    actual_sheet = find_sheet_name(xls, desired_sheet_name)
    if not actual_sheet:
        return pd.DataFrame()

    raw = pd.read_excel(xls, sheet_name=actual_sheet, header=None, engine="openpyxl")
    if raw.empty:
        return pd.DataFrame()

    header_row = None

    for idx, row in raw.iterrows():
        values = [normalize_text(v) for v in row.tolist()]
        if DISTRICT_COL in values:
            header_row = idx
            break

    if header_row is None:
        return pd.DataFrame()

    headers = [normalize_text(v) for v in raw.iloc[header_row].tolist()]
    data = raw.iloc[header_row + 1:].copy()
    data.columns = headers
    data = clean_columns(data)
    data = data.dropna(how="all")

    if DISTRICT_COL in data.columns:
        data = data[data[DISTRICT_COL].notna()]
        data = data[data[DISTRICT_COL].astype(str).str.strip() != ""]

    return data.reset_index(drop=True)


def find_col(df, exact_names=None, contains_all=None):
    if df.empty:
        return None

    exact_names = exact_names or []
    contains_all = contains_all or []

    normalized_map = {normalize_key(col): col for col in df.columns}

    for name in exact_names:
        key_name = normalize_key(name)
        if key_name in normalized_map:
            return normalized_map[key_name]

    if contains_all:
        for col in df.columns:
            col_key = normalize_key(col)
            if all(term.lower() in col_key for term in contains_all):
                return col

    return None


def get_value(row_or_dict, col_name, default=""):
    if row_or_dict is None:
        return default

    if isinstance(row_or_dict, dict):
        for key_name, value in row_or_dict.items():
            if normalize_key(key_name) == normalize_key(col_name):
                return value
        return default

    try:
        for key_name in row_or_dict.index:
            if normalize_key(key_name) == normalize_key(col_name):
                return row_or_dict.get(key_name, default)
    except Exception:
        return default

    return default


def format_number(value, decimals=2):
    if normalize_text(value) == "":
        return ""
    try:
        return f"{float(value):.{decimals}f}"
    except Exception:
        return normalize_text(value)


def format_enrollment(value):
    if normalize_text(value) == "":
        return ""
    try:
        return f"{int(float(value)):,}"
    except Exception:
        return normalize_text(value)


def format_percent(value, decimals=0):
    if normalize_text(value) == "":
        return ""

    try:
        number = float(value)
        pct = number * 100 if abs(number) <= 1 else number
        if decimals == 0:
            return f"{pct:.0f}%"
        return f"{pct:.{decimals}f}%"
    except Exception:
        return normalize_text(value)


def parse_float(value):
    try:
        text = normalize_text(value).replace(",", "")
        if not text:
            return None
        return float(text)
    except Exception:
        return None


def truncate_text(text, max_chars=230):
    text = normalize_text(text)
    if len(text) <= max_chars:
        return text
    cut = text[:max_chars].rsplit(" ", 1)[0]
    return cut + "..."


def export_timestamp():
    return datetime.now().strftime("%B %d, %Y %I:%M %p")


def clean_bullet_text(text):
    text = normalize_text(text)
    text = text.replace("&amp;", "&")
    return text


def compact_items(items, limit=6, max_chars=230):
    output = []
    for item in items:
        item = clean_bullet_text(item)
        if not item:
            continue
        output.append(truncate_text(item, max_chars))
        if len(output) >= limit:
            break
    return output


def priority_rank(priority):
    return {
        "Very High": 4,
        "High": 3,
        "Medium-High": 2,
        "Medium": 1,
    }.get(priority, 0)


def score_sort_value(card):
    return parse_float(card.get("score")) or 0


def enrollment_sort_value(card):
    return parse_float(str(card.get("enrollment", "")).replace(",", "")) or 0


# ============================================================
# TAGS / BADGES
# ============================================================

def tag_class(tag):
    mapping = {
        "Math": "tag-math",
        "MTSS": "tag-mtss",
        "SPED/ELL": "tag-spedell",
        "CCMR": "tag-ccmr",
        "Teacher Capacity": "tag-teacher",
        "Curriculum / HQIM": "tag-hqim",
        "Funding / Grants": "tag-funding",
        "Existing Relationship": "tag-relationship",
        "Literacy / Dyslexia": "tag-literacy",
        "Coaching / Leadership": "tag-coaching",
        "Data Systems": "tag-data",
    }
    return mapping.get(tag, "tag-default")


def badge_html(tag):
    return f'<span class="badge {tag_class(tag)}">{safe_html(tag)}</span>'


def chip_html(item):
    return f'<span class="chip">{safe_html(item)}</span>'


# ============================================================
# PCG EDUCATION SERVICE INTELLIGENCE LIBRARY
# ============================================================

PCG_SERVICE_LIBRARY = [
    {
        "category": "Math Instructional Resources",
        "solution": "Elevation Station Math Games",
        "tags": ["Math", "MTSS", "Teacher Capacity"],
        "keywords": [
            "math fluency", "numeracy", "staAR math", "algebra readiness", "math practice",
            "small group", "intervention", "tier 2", "tutoring", "summer learning",
            "conceptual understanding", "procedural fluency", "student engagement"
        ],
        "buyer_personas": [
            "Math Director", "Curriculum Director", "Elementary Director",
            "Middle School Director", "Intervention Director", "Chief Academic Officer"
        ],
        "positioning": "Position as flexible, standards-aligned math practice that builds fluency through reasoning, discourse, strategy selection, and joyful engagement without replacing the core curriculum.",
        "fit": "Best when districts need classroom-ready math practice, fluency, intervention, enrichment, or extended learning supports.",
        "discovery": [
            "Where do students need more frequent opportunities to practice and explain mathematical thinking?",
            "How are teachers currently supporting fluency without relying only on rote practice?",
            "Where could small-group or station-based resources help teachers respond to student needs?"
        ],
    },
    {
        "category": "K–5 Math Intervention",
        "solution": "Elevation Lessons",
        "tags": ["Math", "MTSS", "Curriculum / HQIM", "Teacher Capacity"],
        "keywords": [
            "k-5 math", "intervention", "remediation", "small group", "enrichment",
            "problem solving", "reasoning", "discourse", "formative assessment",
            "pre assessment", "post assessment", "summer school", "tutoring"
        ],
        "buyer_personas": [
            "Elementary Curriculum Director", "Math Director", "Intervention Director",
            "Chief Academic Officer", "School Improvement Leader"
        ],
        "positioning": "Position as classroom-ready K–5 supplemental lessons that complement core instruction while supporting problem-solving, reasoning, discourse, formative assessment, and targeted intervention.",
        "fit": "Best when districts need K–5 math supports that fit alongside adopted materials and provide teacher-ready routines.",
        "discovery": [
            "Where are K–5 students showing unfinished learning or gaps in foundational mathematics?",
            "How are teachers currently selecting targeted lessons for small-group support?",
            "What evidence helps teachers decide when students are ready to move forward?"
        ],
    },
    {
        "category": "Career-Connected Math / CCMR",
        "solution": "RISE Career & Math Mini Lessons",
        "tags": ["Math", "CCMR", "Teacher Capacity"],
        "keywords": [
            "career readiness", "ccmr", "career pathways", "cte", "industry-based certifications",
            "real world math", "applied math", "middle grades", "career awareness",
            "pathways", "college career military readiness"
        ],
        "buyer_personas": [
            "CCMR Director", "CTE Director", "Math Director",
            "Middle Grades Director", "Curriculum Director"
        ],
        "positioning": "Use when districts want students to connect grade-level math to authentic career contexts, pathway awareness, and future readiness.",
        "fit": "Best when strategic plans connect math, engagement, career pathways, and readiness.",
        "discovery": [
            "Where do students first connect academic skills to future pathways?",
            "How are middle-grade students building awareness of careers that use mathematics?",
            "What would make career-connected math useful without disrupting core pacing?"
        ],
    },
    {
        "category": "Professional Learning Infrastructure",
        "solution": "PCG Playbook Professional Learning Platform",
        "tags": ["Teacher Capacity", "Coaching / Leadership", "SPED/ELL", "Curriculum / HQIM", "Data Systems"],
        "keywords": [
            "professional learning", "asynchronous", "coaching", "mentoring", "onboarding",
            "teacher certification", "in-service credits", "pd management", "learning paths",
            "role-based training", "teacher capacity", "leadership development",
            "dyslexia playbook", "special education playbook", "science of reading"
        ],
        "buyer_personas": [
            "Professional Learning Director", "Chief Academic Officer", "HR/Talent Leader",
            "Special Education Director", "Curriculum Director", "Instructional Coaching Lead", "IT Leader"
        ],
        "positioning": "Position Playbook as the districtwide infrastructure for delivering, managing, personalizing, tracking, and sustaining professional learning and coaching.",
        "fit": "Best when districts need scalable adult-learning infrastructure, role-based learning paths, certificates, coaching tools, and progress reporting.",
        "discovery": [
            "How are professional learning, coaching, and completion data currently tracked?",
            "Where does the district need more role-specific or just-in-time support for educators?",
            "What professional learning needs to be scaled without losing consistency?"
        ],
    },
    {
        "category": "Math Professional Learning",
        "solution": "Asynchronous Math Professional Learning + PLC Implementation Support",
        "tags": ["Math", "Teacher Capacity", "Curriculum / HQIM", "Data Systems"],
        "keywords": [
            "math professional learning", "math pd", "plc", "diagnostics", "assessment",
            "survey", "math coaching", "teacher training", "implementation support",
            "middle grades math", "district-owned content", "playbook"
        ],
        "buyer_personas": [
            "Math Director", "Professional Learning Director", "Curriculum Director",
            "Chief Academic Officer", "Assistant Superintendent"
        ],
        "positioning": "Lead with teacher capacity, PLC implementation, diagnostics, and math instructional coherence; use Playbook as the delivery and sustainability platform when appropriate.",
        "fit": "Best when districts need adult learning, PLC tools, math implementation supports, and a multi-year sustainability path.",
        "discovery": [
            "Where do math PLCs need clearer tools or common routines?",
            "What evidence helps leaders know whether math professional learning is changing classroom practice?",
            "Would district-owned math PL content support long-term sustainability?"
        ],
    },
    {
        "category": "MTSS / Intervention Systems",
        "solution": "PCG MTSS Consulting and Intervention Design",
        "tags": ["MTSS", "Math", "SPED/ELL", "Data Systems"],
        "keywords": [
            "mtss", "tier 2", "tier 3", "intervention", "progress monitoring",
            "data cycles", "student support", "campus variation", "early warning",
            "response to intervention", "rti", "behavior supports"
        ],
        "buyer_personas": [
            "Chief Academic Officer", "Student Support Director", "MTSS Coordinator",
            "Intervention Director", "Special Education Director"
        ],
        "positioning": "Lead with intervention consistency, data-to-action routines, progress monitoring, and campus implementation fidelity.",
        "fit": "Best when districts show intervention, student support, progress monitoring, or campus-variation signals.",
        "discovery": [
            "What tends to slow down the response after students are identified for support?",
            "Where do Tier 2 and Tier 3 routines vary most across campuses?",
            "How do leaders know whether interventions are implemented with fidelity?"
        ],
    },
    {
        "category": "Special Education Systems Improvement",
        "solution": "SPED Systems Redesign, Compliance Monitoring, and Implementation Coherence",
        "tags": ["SPED/ELL", "Teacher Capacity", "Data Systems", "Coaching / Leadership"],
        "keywords": [
            "special education", "sped", "exceptional student services", "compliance",
            "dispute resolution", "iep", "evaluation", "eligibility", "transition planning",
            "behavior supports", "process mapping", "raci", "dashboard", "service delivery",
            "inclusive practices", "co-teaching", "progress monitoring"
        ],
        "buyer_personas": [
            "Special Education Director", "Executive Director of Exceptional Student Services",
            "Chief Academic Officer", "Student Support Leader", "Compliance Director"
        ],
        "positioning": "Frame PCG as a partner for moving special education from compliance-driven activity to coherent, accountable, student-centered implementation through process design, data transparency, ownership structures, and professional learning.",
        "fit": "Best when districts need special education system coherence, compliance visibility, service-delivery improvement, or cross-campus consistency.",
        "discovery": [
            "Where do special education workflows break down across schools?",
            "How are service delivery, compliance patterns, and student progress monitored in real time?",
            "How clear are ownership and decision rights across central office and campus teams?"
        ],
    },
    {
        "category": "Literacy / Dyslexia / Audit Services",
        "solution": "Structured Literacy and Dyslexia Program Review",
        "tags": ["Literacy / Dyslexia", "SPED/ELL", "Teacher Capacity", "Curriculum / HQIM", "Data Systems"],
        "keywords": [
            "literacy", "science of reading", "structured literacy", "dyslexia", "reading disabilities",
            "early literacy", "adolescent literacy", "curriculum audit", "literacy audit",
            "instructional materials review", "evidence-aligned", "reading intervention",
            "phonics", "fluency", "comprehension", "foundational skills"
        ],
        "buyer_personas": [
            "Literacy Director", "Chief Academic Officer", "Curriculum Director",
            "Special Education Director", "State Literacy Lead", "Professional Learning Director"
        ],
        "positioning": "Position PCG as a partner that can review, strengthen, and operationalize evidence-aligned literacy and dyslexia systems through rubrics, audit tools, data analysis, professional learning, and actionable recommendations.",
        "fit": "Best when districts show literacy, dyslexia, reading disability, curriculum review, or evidence-aligned instruction signals.",
        "discovery": [
            "What evidence shows whether literacy instruction is consistently aligned to reading science?",
            "How are dyslexia identification and intervention expectations implemented across campuses?",
            "Would a literacy or dyslexia implementation review help clarify the highest-leverage next steps?"
        ],
    },
    {
        "category": "Coaching Systems / Leadership Development",
        "solution": "Coaching Infrastructure and Leadership Development",
        "tags": ["Coaching / Leadership", "Teacher Capacity", "Data Systems", "Curriculum / HQIM"],
        "keywords": [
            "coaching", "instructional coach", "coaching coordinator", "leadership development",
            "observation", "feedback", "walkthrough", "coach selection", "onboarding",
            "rubric", "look-fors", "implementation fidelity", "leadership pipeline",
            "principal coaching", "coaches of coaches"
        ],
        "buyer_personas": [
            "Professional Learning Director", "Coaching Director", "Chief Academic Officer",
            "School Improvement Director", "HR/Talent Leader", "Literacy Director", "Math Director"
        ],
        "positioning": "Position PCG as a partner that helps districts build coaching infrastructure: observation tools, feedback routines, coach selection/onboarding systems, workshop series, technical assistance, and data-informed continuous improvement.",
        "fit": "Best when districts mention coaching, leadership pipelines, walkthroughs, feedback quality, or implementation monitoring.",
        "discovery": [
            "How does the district define effective coaching practice?",
            "What tools help leaders observe coaching quality and provide actionable feedback?",
            "How are coaches selected, onboarded, and supported after they begin?"
        ],
    },
    {
        "category": "Curriculum / HQIM Implementation",
        "solution": "Curriculum Implementation, PLC Routines, and Change Management",
        "tags": ["Curriculum / HQIM", "Teacher Capacity", "Coaching / Leadership", "Data Systems"],
        "keywords": [
            "curriculum implementation", "hqim", "instructional materials", "adoption fidelity",
            "bluebonnet", "eureka", "carnegie", "curriculum audit", "plc routines",
            "instructional coherence", "walkthrough", "change management"
        ],
        "buyer_personas": [
            "Chief Academic Officer", "Curriculum Director", "Professional Learning Director",
            "School Improvement Director", "Instructional Materials Lead"
        ],
        "positioning": "Frame PCG as an implementation partner helping leaders make adopted materials usable, coherent, and consistently enacted across classrooms.",
        "fit": "Best when districts have adopted or are implementing HQIM, curriculum systems, or instructional transformation plans.",
        "discovery": [
            "Where is implementation strongest and where does variability remain?",
            "What support do teachers need after initial training?",
            "How are leaders monitoring curriculum use and instructional quality?"
        ],
    },
    {
        "category": "Financial / Federal Programs",
        "solution": "Funding Alignment, Medicaid Revenue, and Program Administration Support",
        "tags": ["Funding / Grants", "SPED/ELL", "Data Systems"],
        "keywords": [
            "funding", "grants", "title i", "federal programs", "medicaid",
            "revenue", "reimbursement", "compliance", "audit", "cost settlement",
            "school-based medicaid", "budget", "fiscal sustainability"
        ],
        "buyer_personas": [
            "Chief Financial Officer", "Federal Programs Director", "Special Education Director",
            "Medicaid Coordinator", "Operations Leader"
        ],
        "positioning": "Position PCG as a partner that helps districts align implementation priorities to funding streams, improve compliance, and optimize reimbursement where appropriate.",
        "fit": "Best when districts identify grant alignment, fiscal constraints, federal program compliance, or Medicaid/revenue opportunities.",
        "discovery": [
            "Which strategic priorities already have aligned funding and which do not?",
            "Where are compliance or documentation burdens slowing implementation?",
            "Are there reimbursement or federal program opportunities that could sustain this work?"
        ],
    },
    {
        "category": "School Safety / Behavior Supports",
        "solution": "School Safety, Behavior, and Threat Assessment Support",
        "tags": ["MTSS", "Data Systems", "Coaching / Leadership"],
        "keywords": [
            "school safety", "behavior", "sel", "threat assessment", "climate",
            "culture", "student wellness", "discipline", "prevention", "crisis"
        ],
        "buyer_personas": [
            "Student Support Director", "School Safety Director", "Chief Operations Officer",
            "Behavior Support Lead", "MTSS Coordinator"
        ],
        "positioning": "Use when district priorities include safety, behavior, climate, prevention, or student wellness; frame around consistent systems and early intervention.",
        "fit": "Best when districts connect behavior, safety, climate, and student support systems.",
        "discovery": [
            "How are behavior or safety concerns identified and escalated?",
            "Where do schools need clearer prevention and response routines?",
            "What data helps leaders identify patterns before issues escalate?"
        ],
    },
]


# ============================================================
# COLUMN DETECTION
# ============================================================

def detect_scorecard_columns(scorecard_df):
    return {
        "district": find_col(scorecard_df, exact_names=["District Name"]),
        "enrollment": find_col(scorecard_df, exact_names=["Enrollment"]),
        "strategic_score": find_col(
            scorecard_df,
            exact_names=["Strategic Weighted Score"],
            contains_all=["strategic", "weighted", "score"],
        ),
        "overall_score": find_col(
            scorecard_df,
            exact_names=[
                "Overall Weighted Score (Strategic+Contract+Relationship)",
                "Overall Weighted Score (Strategic+Relationship Weighted)",
                "Overall Weighted Score",
            ],
            contains_all=["overall", "weighted", "score"],
        ),
        "tier": find_col(scorecard_df, exact_names=["Tier"]),
        "existing_contracts": find_col(
            scorecard_df,
            exact_names=["Existing Contracts in Region (Yes/No)"],
            contains_all=["existing", "contracts"],
        ),
        "existing_relationships": find_col(
            scorecard_df,
            exact_names=["Existing Relationships"],
            contains_all=["existing", "relationships"],
        ),
    }


# ============================================================
# LOOKUPS
# ============================================================

def is_substantive_strategic_row(row):
    fields_to_check = [
        "Strategic Plan Themes",
        "Math Improvement Mentioned",
        "Math Priority Strength",
        "Intervention Focus",
        "Intervention Focus Details",
        "Teacher Capacity/PD Focus",
        "Teacher Capacity Details",
        "Career Readiness Mentioned",
        "Career Readiness Details",
        "SPED/ELL Improvement Mentioned",
        "SPED/ELL Details",
        "MTSS/Tiered Support Mentioned",
        "MTSS Details",
        "Curriculum Review/Adoption Activity",
        "Curriculum Details",
        "Active Grants (Yes/No)",
        "Grants Details",
        "Sources",
        "Notes",
    ]
    return sum(1 for field in fields_to_check if normalize_text(get_value(row, field))) >= 2


def build_score_lookup(scorecard_df, score_cols):
    lookup = {}
    district_col = score_cols.get("district")

    if scorecard_df.empty or not district_col:
        return lookup

    for _, row in scorecard_df.iterrows():
        district = normalize_text(row.get(district_col, ""))
        if district:
            lookup[district_key(district)] = row.to_dict()

    return lookup


def build_basic_lookup(basic_df):
    lookup = {}

    if basic_df.empty or DISTRICT_COL not in basic_df.columns:
        return lookup

    for _, row in basic_df.iterrows():
        district = normalize_text(row.get(DISTRICT_COL, ""))
        if district:
            lookup[district_key(district)] = row.to_dict()

    return lookup


def build_csi_tsi_counts(csi_df, tsi_df):
    csi_counts = {}
    tsi_counts = {}

    if not csi_df.empty and DISTRICT_COL in csi_df.columns:
        csi_counts = (
            csi_df.groupby(csi_df[DISTRICT_COL].astype(str).str.upper().str.strip())
            .size()
            .to_dict()
        )

    if not tsi_df.empty and DISTRICT_COL in tsi_df.columns:
        tsi_counts = (
            tsi_df.groupby(tsi_df[DISTRICT_COL].astype(str).str.upper().str.strip())
            .size()
            .to_dict()
        )

    return csi_counts, tsi_counts


def determine_priority(tier, score):
    score_num = parse_float(score)

    if tier == "Tier 1":
        return "Very High"

    if tier == "Tier 2":
        return "High"

    if score_num is not None and score_num >= 3.5:
        return "Medium-High"

    return "Medium"


# ============================================================
# CONTACTS
# ============================================================

def format_contact_display(name="", title="", email="", position=""):
    name = normalize_text(name)
    title = normalize_text(title) or normalize_text(position)
    email = normalize_text(email)

    if name and title and email:
        return f"{name} — {title} | {email}"
    if name and title:
        return f"{name} — {title}"
    if name and email:
        return f"{name} | {email}"
    if title and email:
        return f"{title} | {email}"

    return name or title or email


def build_contacts_lookup_from_contacts(contacts_df):
    lookup = {}

    if contacts_df.empty or DISTRICT_COL not in contacts_df.columns:
        return lookup

    role_priority = [
        "SUPERINTENDENT",
        "CHIEF ACADEMIC OFFICER",
        "ASSISTANT SUPERINTENDENT",
        "CURRICULUM DIRECTOR",
        "MATH COORDINATOR",
        "CTE COLLEGE CAREER READINESS DIRECTOR",
        "SPECIAL EDUCATION DIRECTOR",
        "ESL ELL COORDINATOR",
        "DIRECTOR ASSESSMENT DATA",
        "PROFESSIONAL LEARNING DIRECTOR",
        "ELA LITERACY COORDINATOR",
    ]

    df = contacts_df.copy()

    if "Position" in df.columns:
        df["_priority"] = (
            df["Position"]
            .astype(str)
            .str.upper()
            .apply(lambda x: role_priority.index(x) if x in role_priority else 99)
        )
        df = df.sort_values([DISTRICT_COL, "_priority"])

    for _, row in df.iterrows():
        district = normalize_text(row.get(DISTRICT_COL, ""))
        if not district:
            continue

        first = normalize_text(row.get("First Name", ""))
        last = normalize_text(row.get("Last Name", ""))
        title = normalize_text(row.get("Title", ""))
        position = normalize_text(row.get("Position", ""))
        email = normalize_text(row.get("Email", ""))

        name = f"{first} {last}".strip()
        contact = format_contact_display(
            name=name,
            title=title,
            email=email,
            position=position,
        )

        if contact:
            key = district_key(district)
            lookup.setdefault(key, [])
            if contact not in lookup[key]:
                lookup[key].append(contact)

    return lookup


def build_contacts_lookup_from_leadership(leadership_df):
    lookup = {}

    if leadership_df.empty or DISTRICT_COL not in leadership_df.columns:
        return lookup

    for _, row in leadership_df.iterrows():
        district = normalize_text(row.get(DISTRICT_COL, ""))
        if not district:
            continue

        contacts = []

        superintendent = normalize_text(get_value(row, "Superintendent"))
        superintendent_email = normalize_text(get_value(row, "Email"))

        curriculum_lead = normalize_text(get_value(row, "Curriculum Lead"))
        curriculum_title = normalize_text(get_value(row, "Title")) or "Curriculum / Academic Lead"
        curriculum_email = normalize_text(get_value(row, "Email.1"))

        cte_lead = normalize_text(get_value(row, "CTE Lead"))
        cte_title = normalize_text(get_value(row, "Title.1")) or "CTE / Career Readiness Lead"
        cte_email = normalize_text(get_value(row, "Email.2"))

        math_lead = normalize_text(get_value(row, "Math Lead"))
        math_title = normalize_text(get_value(row, "Title.2")) or "Math Lead"
        math_email = normalize_text(get_value(row, "Email.3"))

        for contact in [
            format_contact_display(superintendent, "Superintendent", superintendent_email),
            format_contact_display(curriculum_lead, curriculum_title, curriculum_email),
            format_contact_display(cte_lead, cte_title, cte_email),
            format_contact_display(math_lead, math_title, math_email),
        ]:
            if contact:
                contact = contact.replace("[", "").replace("]", "")
                contact = contact.replace("(mailto:", " | ").replace(")", "")
                contacts.append(contact)

        lookup[district_key(district)] = contacts[:8]

    return lookup


def get_contacts(district_name, contacts_lookup, leadership_lookup):
    key_name = district_key(district_name)

    if key_name in contacts_lookup and contacts_lookup[key_name]:
        return contacts_lookup[key_name][:8]

    return leadership_lookup.get(key_name, [])[:8]


# ============================================================
# STRATEGIC INTELLIGENCE
# ============================================================

def infer_tags(strategic_row, score_row):
    tags = []
    combined_text = " ".join(
        normalize_text(v)
        for v in strategic_row.to_dict().values()
    ).lower()

    if normalize_text(get_value(strategic_row, "Math Improvement Mentioned")) or normalize_text(get_value(strategic_row, "Math Priority Strength")) or "math" in combined_text:
        tags.append("Math")

    if normalize_text(get_value(strategic_row, "Intervention Focus")) or normalize_text(get_value(strategic_row, "MTSS/Tiered Support Mentioned")) or "mtss" in combined_text or "tier 2" in combined_text or "tier 3" in combined_text:
        tags.append("MTSS")

    if (
        "sped" in combined_text
        or "special education" in combined_text
        or "english learner" in combined_text
        or "emergent bilingual" in combined_text
        or "ell" in combined_text
        or "multilingual" in combined_text
        or "students with disabilities" in combined_text
    ):
        tags.append("SPED/ELL")

    if normalize_text(get_value(strategic_row, "Career Readiness Mentioned")) or normalize_text(get_value(strategic_row, "Career Readiness Details")) or "ccmr" in combined_text or "career" in combined_text:
        tags.append("CCMR")

    if normalize_text(get_value(strategic_row, "Teacher Capacity/PD Focus")) or normalize_text(get_value(strategic_row, "Teacher Capacity Details")) or "professional learning" in combined_text or "coaching" in combined_text or "teacher capacity" in combined_text:
        tags.append("Teacher Capacity")

    if normalize_text(get_value(strategic_row, "Curriculum Review/Adoption Activity")) or normalize_text(get_value(strategic_row, "Curriculum Details")) or "curriculum" in combined_text or "hqim" in combined_text or "instructional materials" in combined_text:
        tags.append("Curriculum / HQIM")

    if normalize_text(get_value(strategic_row, "Active Grants (Yes/No)")) or normalize_text(get_value(strategic_row, "Grants Details")) or "grant" in combined_text or "title i" in combined_text or "funding" in combined_text:
        tags.append("Funding / Grants")

    if "literacy" in combined_text or "reading" in combined_text or "dyslexia" in combined_text or "science of reading" in combined_text or "structured literacy" in combined_text:
        tags.append("Literacy / Dyslexia")

    if "leadership" in combined_text or "walkthrough" in combined_text or "instructional coach" in combined_text or "coaching" in combined_text or "principal" in combined_text:
        tags.append("Coaching / Leadership")

    if "dashboard" in combined_text or "data system" in combined_text or "progress monitoring" in combined_text or "analytics" in combined_text:
        tags.append("Data Systems")

    if normalize_text(get_value(score_row, "Existing Relationships")).lower() == "yes":
        tags.append("Existing Relationship")

    if not tags:
        tags.append("Strategic Review")

    unique = []
    for tag in tags:
        if tag not in unique:
            unique.append(tag)

    return unique


def build_alignment(tags):
    alignment = []

    if "Math" in tags:
        alignment.append("Elevation Station Math Games — K–8 math practice, fluency, reinforcement, and engagement.")
        alignment.append("Elevation intervention curriculum / Elevation Lessons — targeted K–5 support where Tier 2/Tier 3 math or foundational gaps are present.")
        alignment.append("Math professional learning and PLC implementation support — builds teacher capacity and instructional coherence.")

    if "MTSS" in tags:
        alignment.append("PCG MTSS consulting — system design, campus implementation, intervention fidelity, and data-to-action routines.")

    if "SPED/ELL" in tags:
        alignment.append("PCG SPED / multilingual learner support — inclusive practice, service delivery, access to grade-level instruction, compliance, and subgroup progress monitoring.")

    if "CCMR" in tags:
        alignment.append("RISE Career & Math Mini Lessons — grades 6–9 career-connected math, pathway awareness, and applied readiness.")

    if "Curriculum / HQIM" in tags:
        alignment.append("PCG curriculum/HQIM implementation support — adoption fidelity, coaching, PLC routines, and change management.")

    if "Teacher Capacity" in tags:
        alignment.append("PCG professional learning and instructional implementation support — teacher capacity, coaching, data use, and scalable instructional routines.")

    if "Literacy / Dyslexia" in tags:
        alignment.append("PCG structured literacy, dyslexia, and Science of Reading support — program review, professional learning, audit tools, and implementation guidance.")

    if "Coaching / Leadership" in tags:
        alignment.append("PCG coaching infrastructure and leadership development — observation tools, feedback routines, coach selection/onboarding, and leadership capacity building.")

    if "Data Systems" in tags:
        alignment.append("Data transparency and implementation monitoring — dashboards, progress monitoring routines, reporting, and continuous improvement tools.")

    if "Funding / Grants" in tags:
        alignment.append("Funding alignment support — connect implementation supports to existing grant, Title, Medicaid, or strategic funding streams where appropriate.")

    if not alignment:
        alignment.append("Discovery needed — validate strategic needs, stakeholder priorities, and fit before positioning specific resources.")

    return alignment


def build_listen_for(tags):
    tag_map = {
        "Math": ["math growth", "early numeracy", "Algebra readiness", "STAAR math", "student practice", "fluency"],
        "MTSS": ["intervention fidelity", "Tier 2", "Tier 3", "progress monitoring", "campus variation"],
        "SPED/ELL": ["access to grade-level instruction", "service delivery", "emergent bilingual students", "students with disabilities", "subgroup gaps", "compliance"],
        "CCMR": ["pathways", "industry-based certifications", "TSIA2", "dual credit", "career awareness"],
        "Teacher Capacity": ["teacher burden", "coaching", "professional learning", "PLC routines", "instructional consistency"],
        "Curriculum / HQIM": ["adoption fidelity", "HQIM", "curriculum implementation", "Eureka", "Bluebonnet", "instructional materials"],
        "Funding / Grants": ["Title funding", "grant alignment", "LASSO", "federal funds", "implementation funding"],
        "Existing Relationship": ["existing relationship", "current contract", "expansion", "trusted partner"],
        "Literacy / Dyslexia": ["science of reading", "structured literacy", "dyslexia", "foundational skills", "reading intervention"],
        "Coaching / Leadership": ["walkthroughs", "observation tools", "feedback routines", "leadership pipeline", "coach onboarding"],
        "Data Systems": ["dashboards", "analytics", "data cycles", "implementation monitoring", "reporting"],
    }

    listen_for = []

    for tag in tags:
        listen_for.extend(tag_map.get(tag, []))

    listen_for.extend(["implementation barriers", "capacity constraints", "pilot readiness"])

    unique = []
    for item in listen_for:
        if item not in unique:
            unique.append(item)

    return unique


def build_avoid(tags):
    avoid = [
        "Do not lead with a product pitch.",
        "Avoid positioning support as a replacement for the district’s current strategy or adopted curriculum.",
    ]

    if "Curriculum / HQIM" in tags:
        avoid.append("Frame support around implementation and teacher usability rather than another curriculum.")

    if "SPED/ELL" in tags:
        avoid.append("Avoid treating subgroup performance as a side issue; connect it to access, service delivery, and implementation.")

    if "Existing Relationship" in tags:
        avoid.append("Build from existing relationship context before introducing a new idea.")

    if "Literacy / Dyslexia" in tags:
        avoid.append("Avoid implying the district lacks literacy knowledge; position PCG around implementation evidence, coherence, and support.")

    if "Coaching / Leadership" in tags:
        avoid.append("Avoid generic leadership-development language; ask how coaching quality and implementation are currently observed and supported.")

    return avoid


def get_priority_need_opportunity(tags):
    items = []

    if "Math" in tags:
        items.append("Priority: math growth. Need: consistent practice, intervention, and teacher-ready supports. Opportunity: Emerald resources, math professional learning, and PLC implementation.")

    if "Literacy / Dyslexia" in tags:
        items.append("Priority: literacy and dyslexia support. Need: evidence-aligned instruction and implementation coherence. Opportunity: PCG structured literacy review, Science of Reading professional learning, and Playbook.")

    if "MTSS" in tags:
        items.append("Priority: intervention. Need: clear Tier 2/Tier 3 routines. Opportunity: PCG MTSS support, intervention design, and progress-monitoring routines.")

    if "SPED/ELL" in tags:
        items.append("Priority: subgroup access. Need: grade-level expectations with usable supports. Opportunity: PCG SPED/ML systems improvement, inclusive practice, and fidelity monitoring.")

    if "Coaching / Leadership" in tags:
        items.append("Priority: leadership and coaching capacity. Need: consistent observation, feedback, and support routines. Opportunity: PCG coaching infrastructure and leadership development.")

    if "CCMR" in tags:
        items.append("Priority: readiness pathways. Need: earlier career-connected relevance. Opportunity: RISE Career & Math Mini Lessons and CCMR strategy support.")

    if "Curriculum / HQIM" in tags or "Teacher Capacity" in tags:
        items.append("Priority: implementation quality. Need: teacher-ready routines after training. Opportunity: PCG implementation support, Playbook, coaching, and Emerald practice tools.")

    if "Funding / Grants" in tags:
        items.append("Priority: sustainability. Need: funding alignment. Opportunity: connect implementation supports to Title, grant, Medicaid, or strategic funds.")

    if not items:
        items.append("Priority: validate fit. Need: clarify district pain points and implementation constraints. Opportunity: discovery conversation.")

    return items[:4]


def build_relationship_insights(card):
    tags = card.get("tags", [])
    insights = []

    if "Existing Relationship" in tags:
        insights.append("Build from existing relationship credibility before introducing new support.")
    else:
        insights.append("Start with district priorities and implementation realities before discussing resources.")

    if "Curriculum / HQIM" in tags:
        insights.append("Avoid positioning as a competing curriculum; focus on adoption support and classroom usability.")
    elif "Teacher Capacity" in tags:
        insights.append("Emphasize reducing teacher burden and strengthening routines already expected by the district.")
    else:
        insights.append("Use consultative language around practical implementation and measurable proof points.")

    insights.append("Best next step is a narrow proof point tied to a campus, grade band, educator group, or student population.")

    return insights


def build_compact_summary(card):
    tags = card.get("tags", [])
    signals = card.get("signals", [])
    top_tags = [t for t in tags if t not in ["Existing Relationship", "Funding / Grants", "Data Systems"]]

    if card.get("priority") in ["Very High", "High"] and top_tags:
        why = f"{card.get('priority')} priority account with signals around {', '.join(top_tags[:3]).lower()}."
    elif top_tags:
        why = f"High-interest fit around {', '.join(top_tags[:3]).lower()}."
    else:
        why = "Strategic fit should be validated through discovery."

    entry = get_priority_need_opportunity(tags)[0]

    barrier = "Likely barrier: campus variation, teacher capacity, or implementation consistency."

    if "Curriculum / HQIM" in tags:
        barrier = "Likely barrier: turning curriculum or HQIM expectations into consistent classroom routines."
    elif "MTSS" in tags:
        barrier = "Likely barrier: making intervention routines consistent across campuses."
    elif "SPED/ELL" in tags:
        barrier = "Likely barrier: maintaining grade-level access while differentiating support."
    elif "Literacy / Dyslexia" in tags:
        barrier = "Likely barrier: ensuring structured literacy or dyslexia practices are implemented consistently across schools."
    elif "Coaching / Leadership" in tags:
        barrier = "Likely barrier: defining and monitoring coaching quality consistently across campuses."

    return {
        "why": why,
        "entry": entry,
        "barrier": barrier,
        "top_signals": [truncate_text(s, 220) for s in signals[:3]],
        "positioning": get_priority_need_opportunity(tags),
        "relationship": build_relationship_insights(card),
    }


def build_lead_with(card):
    tags = card.get("tags", [])

    if "Existing Relationship" in tags and card.get("tier") == "Tier 1":
        relationship_phrase = "given the existing relationship and strong strategic fit"
    elif "Existing Relationship" in tags:
        relationship_phrase = "building from the existing relationship"
    else:
        relationship_phrase = "as an initial consultative entry point"

    top_tags = [
        tag for tag in tags
        if tag not in ["Existing Relationship", "Funding / Grants", "Data Systems"]
    ]
    top_tags_text = ", ".join(top_tags[:3]).lower() or "strategic implementation support"

    return f"Lead with {top_tags_text} {relationship_phrase}."


def build_conversation_starter(card):
    tags = card.get("tags", [])
    district = card.get("name", "the district")

    if "Math" in tags and "MTSS" in tags:
        return f"For {district}, start with: “As you work on math growth and intervention consistency, where do campuses need the most support turning data into daily instructional action?”"

    if "Literacy / Dyslexia" in tags:
        return f"For {district}, start with: “Where are literacy or dyslexia practices strongest today, and where does implementation still vary across campuses?”"

    if "Coaching / Leadership" in tags:
        return f"For {district}, start with: “How are leaders currently monitoring coaching quality and whether professional learning is changing classroom practice?”"

    if "Curriculum / HQIM" in tags:
        return f"For {district}, start with: “As adopted materials and instructional expectations move into classrooms, where are teachers needing the most practical implementation support?”"

    if "SPED/ELL" in tags:
        return f"For {district}, start with: “Where are multilingual learners or students with disabilities needing more consistent access to grade-level expectations across campuses?”"

    if "CCMR" in tags:
        return f"For {district}, start with: “Where are students first connecting academic skills to pathways, career readiness, and future opportunities?”"

    if "Math" in tags:
        return f"For {district}, start with: “How are campuses currently using math data to decide which students need more practice, intervention, or targeted support?”"

    return f"For {district}, start with: “What district priorities are most difficult to translate into consistent campus-level practice right now?”"


def build_refined_questions(card):
    tags = card.get("tags", [])
    questions = []

    if "Math" in tags:
        questions.append("How are campuses currently using assessment data to decide which students need additional math support?")
    else:
        questions.append("How are campuses currently translating district priorities into daily instructional routines?")

    questions.append("Where does implementation tend to vary most across campuses, grade levels, or student groups?")

    if "MTSS" in tags:
        questions.append("What tends to slow down the response after students are identified for additional support?")

    if "SPED/ELL" in tags:
        questions.append("Where do students with disabilities or multilingual learners need more consistent access to grade-level expectations?")

    if "Literacy / Dyslexia" in tags:
        questions.append("How are structured literacy, dyslexia identification, and reading intervention practices monitored across campuses?")

    if "Coaching / Leadership" in tags:
        questions.append("How do leaders know whether coaching is changing classroom practice?")

    if "Teacher Capacity" in tags or "Curriculum / HQIM" in tags:
        questions.append("What makes it easier or harder for teachers to use new materials or supports consistently after initial training?")

    if "CCMR" in tags:
        questions.append("Where do students begin connecting academic skills to future pathways and career readiness expectations?")

    questions.append("What evidence would tell district and campus leaders that a support model is working well enough to expand?")
    questions.append("If a small pilot were considered, what would make the pilot credible to teachers and principals?")

    unique = []
    for q in questions:
        if q not in unique:
            unique.append(q)

    return unique[:7]


def build_score_band(score):
    score_num = parse_float(score)

    if score_num is None:
        return "Unknown"

    if score_num >= 4.5:
        return "Exceptional"
    if score_num >= 4.0:
        return "Strong"
    if score_num >= 3.5:
        return "Moderate-Strong"
    if score_num >= 3.0:
        return "Selective"

    return "Monitor"


def build_signal_density(tags, signals):
    meaningful_tags = [t for t in tags if t not in ["Existing Relationship"]]

    if len(meaningful_tags) >= 6 or len(signals) >= 8:
        return "High"

    if len(meaningful_tags) >= 3 or len(signals) >= 4:
        return "Medium"

    return "Low"


def build_score_explanation(card):
    tags = card.get("tags", [])
    explanations = []

    overall_num = parse_float(card.get("score"))
    strategic_num = parse_float(card.get("strategic_score"))

    if card.get("tier") == "Tier 1":
        explanations.append("Tier 1 placement signals strong overall fit and near-term prioritization.")
    elif card.get("tier") == "Tier 2":
        explanations.append("Tier 2 placement indicates a strong opportunity with meaningful strategic alignment.")
    elif card.get("tier"):
        explanations.append(f"{card.get('tier')} placement indicates the district is worth monitoring or pursuing selectively.")

    if overall_num is not None:
        if overall_num >= 4.5:
            explanations.append("Very strong overall weighted score suggests unusually strong fit across multiple factors.")
        elif overall_num >= 4.0:
            explanations.append("Strong overall weighted score suggests the district is above average on strategic attractiveness.")
        elif overall_num >= 3.5:
            explanations.append("Moderate-to-strong overall weighted score suggests selective but real opportunity.")

    if strategic_num is not None and strategic_num >= 4.0:
        explanations.append("Strategic score is especially strong, indicating substantive district priorities align to current positioning.")
    elif strategic_num is not None and strategic_num >= 3.5:
        explanations.append("Strategic score is solid, suggesting there is enough signal to justify a guided conversation.")

    if "Math" in tags:
        explanations.append("Math priority signal contributes meaningfully to relevance and urgency.")
    if "MTSS" in tags:
        explanations.append("MTSS/intervention signal increases district need for implementation support and consistency.")
    if "SPED/ELL" in tags:
        explanations.append("SPED/ELL signal indicates subgroup pressure and differentiated access needs.")
    if "Literacy / Dyslexia" in tags:
        explanations.append("Literacy/dyslexia signal suggests relevance for structured literacy, reading intervention, audit/review, or professional learning support.")
    if "Coaching / Leadership" in tags:
        explanations.append("Coaching/leadership signal suggests relevance for observation tools, feedback routines, leadership development, and implementation monitoring.")
    if "CCMR" in tags:
        explanations.append("CCMR signal strengthens relevance when readiness and pathways are district priorities.")
    if "Curriculum / HQIM" in tags:
        explanations.append("Curriculum/HQIM signal indicates implementation support may be more compelling than a net-new program pitch.")
    if "Teacher Capacity" in tags:
        explanations.append("Teacher capacity signal suggests coaching, routines, and usability matter for adoption.")
    if "Existing Relationship" in tags:
        explanations.append("Existing relationship context improves trust and shortens the path to a follow-up conversation.")

    unique = []
    for item in explanations:
        if item not in unique:
            unique.append(item)

    return unique[:7]


def build_score_breakdown(card):
    tags = card.get("tags", [])
    breakdown = []

    if card.get("tier"):
        breakdown.append(f"Tier signal: {card.get('tier')}")
    if card.get("score"):
        breakdown.append(f"Overall weighted score: {card.get('score')}")
    if card.get("strategic_score"):
        breakdown.append(f"Strategic weighted score: {card.get('strategic_score')}")

    drivers = {
        "Math": "Strategic fit driver: math improvement priority",
        "MTSS": "Strategic fit driver: intervention / MTSS need",
        "SPED/ELL": "Strategic fit driver: subgroup access and performance pressure",
        "Literacy / Dyslexia": "Strategic fit driver: literacy, dyslexia, or reading intervention priority",
        "Coaching / Leadership": "Implementation driver: coaching quality, leadership capacity, or walkthrough routines",
        "CCMR": "Strategic fit driver: college, career, and military readiness",
        "Teacher Capacity": "Implementation driver: teacher capacity or professional learning",
        "Curriculum / HQIM": "Implementation driver: curriculum or HQIM adoption fidelity",
        "Funding / Grants": "Funding driver: grant or aligned funding signal",
        "Existing Relationship": "Relationship driver: existing relationship or contract signal",
        "Data Systems": "Implementation driver: data transparency, dashboards, or monitoring routines",
    }

    for tag, text in drivers.items():
        if tag in tags:
            breakdown.append(text)

    return breakdown


def build_next_moves(card):
    tags = card.get("tags", [])
    next_moves = []

    if "Math" in tags and "MTSS" in tags:
        next_moves.append("Lead with how math goals and intervention routines are being implemented across campuses.")
        next_moves.append("Look for a pilot entry point tied to progress monitoring, small-group support, or campus variation.")
    elif "Math" in tags:
        next_moves.append("Anchor the conversation in current math goals, student practice, and growth monitoring.")
        next_moves.append("Explore whether a short pilot could produce a visible proof point.")

    if "Literacy / Dyslexia" in tags:
        next_moves.append("Explore whether a structured literacy, dyslexia, or reading implementation review would help clarify priority next steps.")

    if "MTSS" in tags:
        next_moves.append("Start with the district’s intervention workflow and where implementation slows down.")

    if "Curriculum / HQIM" in tags:
        next_moves.append("Frame support around implementation fidelity and teacher usability within adopted materials.")

    if "Coaching / Leadership" in tags:
        next_moves.append("Ask how coaching quality, walkthroughs, and feedback routines are currently monitored across campuses.")

    if "SPED/ELL" in tags:
        next_moves.append("Test where subgroup access or differentiated implementation needs are creating friction.")

    if "Existing Relationship" in tags:
        next_moves.append("Build from existing relationship credibility and connect to current district priorities.")
    else:
        next_moves.append("Use discovery to validate urgency before positioning a larger district-wide motion.")

    unique = []
    for item in next_moves:
        if item not in unique:
            unique.append(item)

    return unique[:5]


# ============================================================
# SERVICE MATCHING
# ============================================================

def service_keyword_score(blob, keywords):
    hits = 0
    matched = []

    for kw in keywords:
        kw_norm = normalize_key(kw)
        if kw_norm and kw_norm in blob:
            hits += 1
            matched.append(kw)

    if not keywords:
        return 0, []

    return min(hits / max(len(keywords), 1), 1), matched[:6]


def match_pcg_services_to_card(card, max_matches=5):
    card_tags = set(card.get("tags", []))
    blob = search_blob(card)
    matches = []

    for service in PCG_SERVICE_LIBRARY:
        service_tags = set(service.get("tags", []))
        tag_overlap = card_tags.intersection(service_tags)
        tag_score = len(tag_overlap) / max(len(service_tags), 1)

        keyword_score, keyword_hits = service_keyword_score(blob, service.get("keywords", []))

        priority_boost = priority_rank(card.get("priority")) / 4 * 0.10
        relationship_boost = 0.05 if "Existing Relationship" in card_tags else 0

        confidence = (tag_score * 0.55) + (keyword_score * 0.30) + priority_boost + relationship_boost

        if confidence > 0.12:
            match_reason_bits = []

            if tag_overlap:
                match_reason_bits.append("Matched tags: " + ", ".join(sorted(tag_overlap)))

            if keyword_hits:
                match_reason_bits.append("Matched signals: " + ", ".join(keyword_hits[:4]))

            if not match_reason_bits:
                match_reason_bits.append("General strategic fit based on district profile.")

            if confidence >= 0.65:
                fit_level = "High"
            elif confidence >= 0.38:
                fit_level = "Medium-High"
            else:
                fit_level = "Selective"

            matches.append({
                "category": service.get("category", ""),
                "solution": service.get("solution", ""),
                "positioning": service.get("positioning", ""),
                "fit": service.get("fit", ""),
                "buyer_personas": service.get("buyer_personas", []),
                "discovery": service.get("discovery", []),
                "confidence": confidence,
                "fit_level": fit_level,
                "match_reason": "; ".join(match_reason_bits),
            })

    matches = sorted(matches, key=lambda m: m["confidence"], reverse=True)

    return matches[:max_matches]


# ============================================================
# CARD BUILDING
# ============================================================

def build_card(
    strategic_row,
    score_row,
    basic_lookup,
    contacts_lookup,
    leadership_lookup,
    csi_counts,
    tsi_counts,
    score_cols,
):
    district_name = normalize_text(get_value(strategic_row, DISTRICT_COL))
    district_id = district_key(district_name)

    raw_score = score_row.get(score_cols.get("overall_score"), "")
    raw_strategic_score = score_row.get(score_cols.get("strategic_score"), "") if score_cols.get("strategic_score") else ""

    score = format_number(raw_score, 2)
    strategic_score = format_number(raw_strategic_score, 2)
    tier = normalize_text(score_row.get(score_cols.get("tier"), "")) if score_cols.get("tier") else ""
    enrollment = format_enrollment(score_row.get(score_cols.get("enrollment"), "")) if score_cols.get("enrollment") else ""

    tags = infer_tags(strategic_row, score_row)
    card_priority = determine_priority(tier, raw_score)
    basic = basic_lookup.get(district_id, {})

    signals = []

    themes = normalize_text(get_value(strategic_row, "Strategic Plan Themes"))
    math_strength = normalize_text(get_value(strategic_row, "Math Priority Strength"))
    intervention_details = normalize_text(get_value(strategic_row, "Intervention Focus Details"))
    sped_ell_details = normalize_text(get_value(strategic_row, "SPED/ELL Details"))
    teacher_details = normalize_text(get_value(strategic_row, "Teacher Capacity Details"))
    career_details = normalize_text(get_value(strategic_row, "Career Readiness Details"))
    mtss_details = normalize_text(get_value(strategic_row, "MTSS Details"))
    curriculum_details = normalize_text(get_value(strategic_row, "Curriculum Details"))
    grants_details = normalize_text(get_value(strategic_row, "Grants Details"))
    notes = normalize_text(get_value(strategic_row, "Notes"))

    if themes:
        signals.append(f"Strategic themes: {themes}")

    if math_strength:
        signals.append(f"Math signal: {math_strength}")

    csi = normalize_text(get_value(basic, "CSI Schools")) or str(csi_counts.get(district_id, ""))
    tsi = normalize_text(get_value(basic, "TSI Schools")) or str(tsi_counts.get(district_id, ""))

    if csi or tsi:
        signals.append(f"Accountability pressure: {csi or '0'} CSI schools and {tsi or '0'} TSI schools.")

    context_parts = []
    context_fields = [
        ("schools", "Number of Schools", "text"),
        ("grade span", "Grade Span Served", "text"),
        ("setting", "Urban/Suburban/Rural", "text"),
        ("economically disadvantaged", "% Economically Disadvantaged", "percent"),
        ("English learner", "% English Learner", "percent"),
        ("special education", "% Special Education", "percent"),
        ("major student group", "Major Student Groups", "text"),
        ("student growth trend", "Student Growth Trend", "percent1"),
    ]

    for label, col, kind in context_fields:
        raw_value = get_value(basic, col)
        value = normalize_text(raw_value)

        if value:
            if kind == "percent":
                value = format_percent(raw_value, 0)
            elif kind == "percent1":
                value = format_percent(raw_value, 1)
            context_parts.append(f"{label}: {value}")

    if context_parts:
        signals.append("District context: " + "; ".join(context_parts) + ".")

    if intervention_details:
        signals.append(f"Intervention signal: {intervention_details}")
    if mtss_details:
        signals.append(f"MTSS signal: {mtss_details}")
    if sped_ell_details:
        signals.append(f"SPED/ELL signal: {sped_ell_details}")
    if teacher_details:
        signals.append(f"Teacher capacity signal: {teacher_details}")
    if career_details:
        signals.append(f"CCMR / career readiness signal: {career_details}")
    if curriculum_details:
        signals.append(f"Curriculum / implementation signal: {curriculum_details}")
    if grants_details:
        signals.append(f"Funding / grants signal: {grants_details}")
    if notes:
        signals.append(f"Additional notes: {notes}")

    existing_relationships = normalize_text(get_value(score_row, "Existing Relationships"))
    existing_contracts = normalize_text(get_value(score_row, "Existing Contracts in Region (Yes/No)"))

    if existing_contracts and existing_relationships:
        signals.append(f"Relationship context: existing contracts in region = {existing_contracts}; existing relationships = {existing_relationships}.")
    elif existing_relationships:
        signals.append(f"Relationship context: existing relationships = {existing_relationships}.")
    elif existing_contracts:
        signals.append(f"Relationship context: existing contracts in region = {existing_contracts}.")

    card = {
        "name": district_name,
        "tier": tier,
        "score": score,
        "strategic_score": strategic_score,
        "enrollment": enrollment,
        "priority": card_priority,
        "tags": tags,
        "contacts": get_contacts(district_name, contacts_lookup, leadership_lookup),
        "signals": signals,
        "alignment": build_alignment(tags),
        "listen": build_listen_for(tags),
        "avoid": build_avoid(tags),
    }

    card["score_band"] = build_score_band(card.get("score"))
    card["signal_density"] = build_signal_density(tags, signals)
    card["lead"] = build_lead_with(card)
    card["questions"] = build_refined_questions(card)
    card["compact"] = build_compact_summary(card)
    card["score_explanation"] = build_score_explanation(card)
    card["score_breakdown"] = build_score_breakdown(card)
    card["next_moves"] = build_next_moves(card)
    card["conversation_starter"] = build_conversation_starter(card)
    card["pcg_service_matches"] = match_pcg_services_to_card(card, max_matches=6)

    return card


# ============================================================
# WORKBOOK LOADING WITH CACHING
# ============================================================

def get_workbook_source(uploaded_file):
    if uploaded_file is not None:
        return uploaded_file.getvalue(), "Uploaded workbook override", "bytes"

    if os.path.exists(BUILT_IN_WORKBOOK):
        return BUILT_IN_WORKBOOK, "Built-in workbook", "path"

    return None, "No workbook found", "none"


@st.cache_data(show_spinner=False)
def load_cards_from_workbook_cached(workbook_payload, source_type):
    if source_type == "bytes":
        xls = pd.ExcelFile(BytesIO(workbook_payload), engine="openpyxl")
    else:
        xls = pd.ExcelFile(workbook_payload, engine="openpyxl")

    strategic_df = read_table_sheet(xls, SHEET_STRATEGIC)
    scorecard_df = read_table_sheet(xls, SHEET_SCORECARD)
    basic_df = read_table_sheet(xls, SHEET_BASIC)
    contacts_df = read_table_sheet(xls, SHEET_CONTACTS)
    leadership_df = read_table_sheet(xls, SHEET_LEADERSHIP)
    csi_df = read_table_sheet(xls, SHEET_CSI)
    tsi_df = read_table_sheet(xls, SHEET_TSI)

    if strategic_df.empty or scorecard_df.empty:
        return (
            [],
            xls.sheet_names,
            strategic_df,
            scorecard_df,
            basic_df,
            contacts_df,
            leadership_df,
            csi_df,
            tsi_df,
            pd.DataFrame(),
            {},
        )

    score_cols = detect_scorecard_columns(scorecard_df)
    score_lookup = build_score_lookup(scorecard_df, score_cols)
    basic_lookup = build_basic_lookup(basic_df)
    contacts_lookup = build_contacts_lookup_from_contacts(contacts_df)
    leadership_lookup = build_contacts_lookup_from_leadership(leadership_df)
    csi_counts, tsi_counts = build_csi_tsi_counts(csi_df, tsi_df)

    cards = []
    audit_rows = []

    for _, strategic_row in strategic_df.iterrows():
        district_name = normalize_text(get_value(strategic_row, DISTRICT_COL))

        if not district_name:
            continue

        district_id = district_key(district_name)
        substantive = is_substantive_strategic_row(strategic_row)
        score_row = score_lookup.get(district_id, {})
        matched = bool(score_row)

        tier = normalize_text(score_row.get(score_cols.get("tier"), "")) if matched and score_cols.get("tier") else ""
        overall = score_row.get(score_cols.get("overall_score"), "") if matched and score_cols.get("overall_score") else ""

        eligible = substantive and matched and bool(tier) and bool(normalize_text(overall))

        audit_rows.append({
            "District": district_name,
            "Substantive Strategic Indicators": substantive,
            "Matched Master Scorecard": matched,
            "Detected Tier Column": score_cols.get("tier"),
            "Tier": tier,
            "Detected Overall Score Column": score_cols.get("overall_score"),
            "Overall Score": normalize_text(overall),
            "Eligible": eligible,
        })

        if eligible:
            cards.append(
                build_card(
                    strategic_row,
                    score_row,
                    basic_lookup,
                    contacts_lookup,
                    leadership_lookup,
                    csi_counts,
                    tsi_counts,
                    score_cols,
                )
            )

    audit_df = pd.DataFrame(audit_rows)

    return (
        cards,
        xls.sheet_names,
        strategic_df,
        scorecard_df,
        basic_df,
        contacts_df,
        leadership_df,
        csi_df,
        tsi_df,
        audit_df,
        score_cols,
    )


# ============================================================
# SEARCH / FILTER / SORT
# ============================================================

def search_blob(card):
    values = [
        card.get("name", ""),
        card.get("tier", ""),
        str(card.get("score", "")),
        str(card.get("strategic_score", "")),
        card.get("enrollment", ""),
        card.get("priority", ""),
        card.get("lead", ""),
        card.get("conversation_starter", ""),
        card.get("score_band", ""),
        card.get("signal_density", ""),
    ]

    for field in [
        "tags",
        "signals",
        "alignment",
        "contacts",
        "questions",
        "listen",
        "avoid",
        "score_explanation",
        "score_breakdown",
        "next_moves",
    ]:
        values.extend(card.get(field, []))

    for match in card.get("pcg_service_matches", []):
        values.extend([
            match.get("category", ""),
            match.get("solution", ""),
            match.get("positioning", ""),
            match.get("fit", ""),
            match.get("match_reason", ""),
            " ".join(match.get("buyer_personas", [])),
            " ".join(match.get("discovery", [])),
        ])

    return " ".join(map(str, values)).lower()


def filter_cards(cards, query, selected_tiers, selected_tags, shortlist, show_shortlist_only):
    query = query.strip().lower()
    filtered = []

    for card in cards:
        if selected_tiers and card.get("tier") not in selected_tiers:
            continue

        if selected_tags and not any(tag in card.get("tags", []) for tag in selected_tags):
            continue

        if show_shortlist_only and card.get("name") not in shortlist:
            continue

        if query and query not in search_blob(card):
            continue

        filtered.append(card)

    return filtered


def sort_cards(cards, sort_mode):
    if sort_mode == "Priority then Score":
        return sorted(cards, key=lambda c: (priority_rank(c.get("priority")), score_sort_value(c)), reverse=True)

    if sort_mode == "Overall Score":
        return sorted(cards, key=score_sort_value, reverse=True)

    if sort_mode == "Strategic Score":
        return sorted(cards, key=lambda c: parse_float(c.get("strategic_score")) or 0, reverse=True)

    if sort_mode == "Enrollment":
        return sorted(cards, key=enrollment_sort_value, reverse=True)

    if sort_mode == "Signal Density":
        density_rank = {"High": 3, "Medium": 2, "Low": 1}
        return sorted(cards, key=lambda c: (density_rank.get(c.get("signal_density"), 0), score_sort_value(c)), reverse=True)

    if sort_mode == "District A-Z":
        return sorted(cards, key=lambda c: c.get("name", ""))

    return cards


# ============================================================
# DOCX EXPORT HELPERS
# ============================================================

def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_bullets(cell, items, limit=None):
    use_items = items[:limit] if limit else items

    if not use_items:
        cell.text = ""
        return

    cell.text = ""

    for item in use_items:
        if normalize_text(item):
            p = cell.add_paragraph(style=None)
            p.text = f"• {item}"


def build_district_overview_text(card):
    name = card.get("name", "")
    tier = card.get("tier", "")
    priority = card.get("priority", "")
    enrollment = card.get("enrollment", "")
    score = card.get("score", "")
    strategic_score = card.get("strategic_score", "")

    tags = [
        t for t in card.get("tags", [])
        if t not in ["Existing Relationship", "Funding / Grants", "Data Systems"]
    ]

    tag_phrase = ", ".join(tags[:4]).lower() if tags else "district strategic priorities"

    overview = f"{name} is a {priority.lower() if priority else 'priority'} opportunity with signals around {tag_phrase}. "

    if enrollment:
        overview += f"The district serves approximately {enrollment} students. "

    score_bits = []
    if tier:
        score_bits.append(tier)
    if score:
        score_bits.append(f"overall score {score}")
    if strategic_score:
        score_bits.append(f"strategic score {strategic_score}")

    if score_bits:
        overview += "Current scoring profile: " + ", ".join(score_bits) + ". "

    compact = card.get("compact", {})
    why = clean_bullet_text(compact.get("why", ""))
    barrier = clean_bullet_text(compact.get("barrier", ""))

    if why:
        overview += why + " "
    if barrier:
        overview += barrier

    return overview.strip()


def build_strategic_objectives(card):
    objectives = []
    signals = card.get("signals", [])

    for signal in signals:
        s = clean_bullet_text(signal)

        if s.lower().startswith("strategic themes:"):
            objectives.append(s.replace("Strategic themes:", "").strip())
        elif "Math signal:" in s:
            objectives.append(s.replace("Math signal:", "Improve mathematics outcomes:").strip())
        elif "MTSS signal:" in s or "Intervention signal:" in s:
            cleaned = s.replace("MTSS signal:", "").replace("Intervention signal:", "").strip()
            objectives.append(f"Strengthen intervention systems and implementation consistency: {cleaned}")
        elif "SPED/ELL signal:" in s:
            objectives.append(s.replace("SPED/ELL signal:", "Improve outcomes and access for priority student groups:").strip())
        elif "Teacher capacity signal:" in s:
            objectives.append(s.replace("Teacher capacity signal:", "Build educator capacity and improve instructional consistency:").strip())
        elif "CCMR / career readiness signal:" in s:
            objectives.append(s.replace("CCMR / career readiness signal:", "Increase college, career, and military readiness:").strip())
        elif "Curriculum / implementation signal:" in s:
            objectives.append(s.replace("Curriculum / implementation signal:", "Strengthen curriculum implementation and instructional coherence:").strip())

    if not objectives:
        compact = card.get("compact", {})
        if compact.get("why"):
            objectives.append(clean_bullet_text(compact.get("why")))
        if compact.get("entry"):
            objectives.append(clean_bullet_text(compact.get("entry")))

    return compact_items(objectives, limit=7, max_chars=275)


def build_pain_points(card):
    pain_points = []
    compact = card.get("compact", {})
    barrier = clean_bullet_text(compact.get("barrier", ""))

    if barrier:
        pain_points.append(barrier.replace("Likely barrier:", "").strip())

    tags = card.get("tags", [])

    if "Math" in tags:
        pain_points.append("Math growth, fluency, or grade-level readiness may require more consistent practice, intervention, and progress monitoring.")

    if "Literacy / Dyslexia" in tags:
        pain_points.append("Literacy, dyslexia, or reading intervention practices may need stronger coherence, monitoring, and evidence-aligned implementation.")

    if "MTSS" in tags:
        pain_points.append("Intervention systems may vary across campuses, especially in how students are identified, supported, and monitored.")

    if "SPED/ELL" in tags:
        pain_points.append("Students with disabilities, emergent bilingual students, or other priority groups may need more consistent access to grade-level instruction and supports.")

    if "Teacher Capacity" in tags:
        pain_points.append("Professional learning may not yet be translating into consistent classroom practice across campuses.")

    if "Curriculum / HQIM" in tags:
        pain_points.append("Curriculum or HQIM expectations may require stronger implementation routines, coaching, PLC structures, or fidelity monitoring.")

    if "Coaching / Leadership" in tags:
        pain_points.append("Coaching and leadership routines may lack consistent observation tools, feedback cycles, or implementation monitoring.")

    if "CCMR" in tags:
        pain_points.append("Students may need earlier and more explicit connections between academic skills, career pathways, and readiness expectations.")

    if "Funding / Grants" in tags:
        pain_points.append("The district may need help aligning implementation supports to available grant, Title, Medicaid, or strategic funding streams.")

    return compact_items(pain_points, limit=7, max_chars=245)


def build_pcg_entry_points(card):
    rows = []

    for match in card.get("pcg_service_matches", [])[:6]:
        rows.append((
            match.get("category", ""),
            match.get("solution", ""),
            match.get("fit", "")
        ))

    if not rows:
        for item in card.get("alignment", [])[:5]:
            rows.append(("Strategic Fit", item, "Validate need and fit through discovery."))

    return rows


def build_decision_makers(card):
    contacts = card.get("contacts", [])
    return compact_items(contacts, limit=8, max_chars=180)


def build_nepq_approach(card):
    tags = card.get("tags", [])
    compact = card.get("compact", {})
    approach = []

    lead = clean_bullet_text(card.get("lead", ""))
    if lead:
        approach.append(lead)

    entry = clean_bullet_text(compact.get("entry", ""))
    if entry:
        approach.append(f"Position around this entry point: {entry}")

    top_matches = card.get("pcg_service_matches", [])[:2]
    if top_matches:
        top_solution_text = "; ".join([m.get("solution", "") for m in top_matches])
        approach.append(f"Most relevant PCG solution pathways to validate: {top_solution_text}.")

    approach.append(
        "Use a consultative approach that starts with district priorities, current implementation realities, and desired outcomes before introducing a specific PCG or Emerald solution."
    )

    if "Existing Relationship" in tags:
        approach.append("Because a relationship signal exists, build from known PCG credibility before introducing a new service pathway.")

    questions = []

    if "Math" in tags:
        questions.append("What is keeping math growth from becoming more consistent across campuses or grade levels?")

    if "Literacy / Dyslexia" in tags:
        questions.append("Where are literacy or dyslexia practices strongest today, and where does implementation still vary?")

    if "MTSS" in tags:
        questions.append("When students are identified for support, where does the intervention process slow down or vary most?")

    if "SPED/ELL" in tags:
        questions.append("Where are students with disabilities or emergent bilingual students still not receiving consistent access to grade-level expectations?")

    if "Teacher Capacity" in tags or "Curriculum / HQIM" in tags:
        questions.append("Where is professional learning or curriculum implementation not yet translating into classroom execution?")

    if "Coaching / Leadership" in tags:
        questions.append("How are coaching quality, walkthroughs, and feedback routines monitored today?")

    if "CCMR" in tags:
        questions.append("Where do students begin connecting academic skills to future pathways and career readiness expectations?")

    questions.append("If PCG helped solve one implementation barrier this year, which barrier would create the most visible impact?")
    questions.append("What would make a small pilot credible enough for principals, teachers, and district leaders to support expansion?")

    return compact_items(approach, limit=5, max_chars=275), compact_items(questions, limit=8, max_chars=230)


def build_profile_next_steps(card):
    next_steps = []

    for item in card.get("next_moves", []):
        next_steps.append(clean_bullet_text(item))

    contacts = card.get("contacts", [])
    if contacts:
        next_steps.append("Use the listed decision-makers and influencers to identify the best first conversation owner and confirm current district priorities.")

    matches = card.get("pcg_service_matches", [])
    if matches:
        next_steps.append(f"Validate whether the strongest initial solution pathway is {matches[0].get('solution', 'the top PCG solution pathway')}.")

    next_steps.append("Prepare a focused pilot, diagnostic, or implementation-support concept tied to one measurable district priority, campus group, grade band, or student population.")

    if "Funding / Grants" in card.get("tags", []):
        next_steps.append("Identify aligned funding sources, including Title, grant, Medicaid, or strategic funds, before proposing scope.")

    return compact_items(next_steps, limit=6, max_chars=245)


# ============================================================
# DISTRICT PROFILE WORD EXPORT
# ============================================================

def build_docx(cards, data_source_label="Built-in workbook"):
    doc = Document()

    title = doc.add_heading("District Profile Document", 0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(23, 54, 93)

    doc.add_paragraph(
        "Business development profile generated from the Strategic District Field Guide workbook."
    )
    doc.add_paragraph(f"Generated: {export_timestamp()}")
    doc.add_paragraph(f"Data source: {data_source_label}")

    for idx, card in enumerate(cards):
        doc.add_heading(card.get("name", "District Profile"), level=1)

        profile_line = (
            f"{card.get('tier', '')} | "
            f"Overall Score {card.get('score', '')} | "
            f"Strategic Score {card.get('strategic_score', '')} | "
            f"Enrollment {card.get('enrollment', '')} | "
            f"Priority {card.get('priority', '')}"
        )
        doc.add_paragraph(profile_line)

        doc.add_heading("1. District Overview", level=2)
        doc.add_paragraph(build_district_overview_text(card))

        doc.add_heading("2. Strategic Objectives", level=2)
        strategic_objectives = build_strategic_objectives(card)
        if strategic_objectives:
            for item in strategic_objectives:
                doc.add_paragraph(item, style="List Bullet")
        else:
            doc.add_paragraph("Strategic objectives should be validated through discovery.", style="List Bullet")

        doc.add_heading("3. Pain Points / Problems to Solve", level=2)
        pain_points = build_pain_points(card)
        if pain_points:
            for item in pain_points:
                doc.add_paragraph(item, style="List Bullet")
        else:
            doc.add_paragraph("Pain points should be validated through discovery.", style="List Bullet")

        doc.add_heading("4. PCG Entry Points", level=2)

        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = ["District Goal / Need", "PCG Solution", "Description of Fit"]
        hdr_cells = table.rows[0].cells

        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            shade_cell(hdr_cells[i], "1F4E79")
            for p in hdr_cells[i].paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.bold = True

        for goal, solution, fit in build_pcg_entry_points(card):
            row_cells = table.add_row().cells
            row_cells[0].text = goal
            row_cells[1].text = solution
            row_cells[2].text = fit

            for cell in row_cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

        doc.add_heading("5. Decision-Makers and Influencers", level=2)
        contacts = build_decision_makers(card)
        if contacts:
            for item in contacts:
                doc.add_paragraph(item, style="List Bullet")
        else:
            doc.add_paragraph("No contacts were available in the workbook for this district.", style="List Bullet")

        doc.add_heading("6. Recommended Approach to Leadership with NEPQ Integration", level=2)
        approach_items, nepq_questions = build_nepq_approach(card)

        doc.add_paragraph("Recommended positioning:")
        for item in approach_items:
            doc.add_paragraph(item, style="List Bullet")

        doc.add_paragraph("NEPQ-style discovery questions:")
        for question in nepq_questions:
            doc.add_paragraph(question, style="List Bullet")

        doc.add_heading("7. Next Steps", level=2)
        for item in build_profile_next_steps(card):
            doc.add_paragraph(item, style="List Bullet")

        if idx < len(cards) - 1:
            doc.add_page_break()

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# ============================================================
# MATRIX EXPORT
# ============================================================

def build_matrix_docx(cards, data_source_label="Built-in workbook"):
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)

    title = doc.add_heading("Texas District Strategic Positioning Matrix", 0)

    for run in title.runs:
        run.font.color.rgb = RGBColor(31, 78, 121)

    doc.add_paragraph("Scope: Currently filtered districts from the Strategic District Field Guide.")
    doc.add_paragraph("Designed for deeper review after quick conference scanning in the app.")
    doc.add_paragraph(f"Generated: {export_timestamp()}")
    doc.add_paragraph(f"Data source: {data_source_label}")

    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    headers = [
        "District / Score / Contacts",
        "District Overview",
        "Top PCG Pathways",
        "Strategic Signal Analysis",
        "Discovery Questions",
        "Next Moves",
    ]

    hdr_cells = table.rows[0].cells

    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        shade_cell(hdr_cells[i], "1F4E79")
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.bold = True

    for card in cards:
        compact = card.get("compact", {})
        row_cells = table.add_row().cells

        for cell in row_cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

        shade_cell(row_cells[0], "EFF6FF")

        contacts = card.get("contacts", [])[:4]
        row_cells[0].text = (
            f"{card.get('name','')}\n"
            f"{card.get('tier','')} | Overall {card.get('score','')}\n"
            f"Strategic {card.get('strategic_score','')}\n"
            f"Priority: {card.get('priority','')}\n"
            f"Enrollment: {card.get('enrollment','')}\n"
            f"Contacts: " + "; ".join(contacts)
        )

        row_cells[1].text = ""
        add_bullets(
            row_cells[1],
            [
                compact.get("why", ""),
                compact.get("entry", ""),
                compact.get("barrier", ""),
            ],
            limit=4
        )

        row_cells[2].text = ""
        add_bullets(
            row_cells[2],
            [
                f"{m.get('solution','')} ({m.get('fit_level','')}) — {m.get('fit','')}"
                for m in card.get("pcg_service_matches", [])[:4]
            ],
            limit=4
        )

        row_cells[3].text = ""
        add_bullets(
            row_cells[3],
            card.get("score_explanation", [])[:3] + compact.get("top_signals", [])[:2],
            limit=5
        )

        row_cells[4].text = ""
        add_bullets(row_cells[4], card.get("questions", []), limit=6)

        row_cells[5].text = ""
        add_bullets(row_cells[5], build_profile_next_steps(card), limit=5)

    doc.add_paragraph(
        "Note: This matrix is designed for strategy review and conference/meeting preparation. Validate against current district conversations before final pursuit decisions."
    )

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# ============================================================
# RENDERING
# ============================================================

def render_summary_grid(card):
    compact = card.get("compact", {})

    st.markdown(f"""
    <div class="summary-grid">
        <div class="summary-box">
            <div class="summary-label">Why it matters</div>
            <div class="summary-value">{safe_html(compact.get('why', ''))}</div>
        </div>
        <div class="summary-box">
            <div class="summary-label">Best entry point</div>
            <div class="summary-value">{safe_html(compact.get('entry', ''))}</div>
        </div>
        <div class="summary-box">
            <div class="summary-label">Likely barrier</div>
            <div class="summary-value">{safe_html(compact.get('barrier', ''))}</div>
        </div>
    </div>
    """, unsafe_allow_html=True)


def render_quick_prep(card):
    top_listen = card.get("listen", [])[:8]
    best_question = card.get("questions", [""])[0] if card.get("questions") else ""
    key_guidance = card.get("avoid", [""])[0] if card.get("avoid") else ""
    chips = "".join(chip_html(item) for item in top_listen)

    st.markdown(f"""
    <div class="quickprep">
        <div class="quickprep-title">Quick Prep</div>
        <strong>Conversation starter:</strong> {safe_html(card.get("conversation_starter", ""))}<br><br>
        <strong>Best opening question:</strong> {safe_html(best_question)}<br>
        <strong>Listen for:</strong><br>{chips}<br>
        <strong>Guidance:</strong> {safe_html(key_guidance)}
    </div>
    """, unsafe_allow_html=True)


def render_pcg_solution_matches(card):
    matches = card.get("pcg_service_matches", [])

    st.markdown("**Recommended PCG Solution Pathways**")

    if not matches:
        st.markdown("_No specific PCG solution pathway was identified. Use discovery to validate fit._")
        return

    for m in matches[:5]:
        personas = ", ".join(m.get("buyer_personas", [])[:4])
        discovery = m.get("discovery", [])
        first_question = discovery[0] if discovery else ""

        st.markdown(f"""
        <div class="solution-card">
            <strong>{safe_html(m.get("solution", ""))}</strong>
            <span class="priority priority-medium-high">{safe_html(m.get("fit_level", ""))} Fit</span><br>
            <span class="helper-text">{safe_html(m.get("category", ""))}</span><br><br>
            <strong>Why it matches:</strong> {safe_html(m.get("match_reason", ""))}<br>
            <strong>How to position:</strong> {safe_html(m.get("positioning", ""))}<br>
            <strong>Likely buyers:</strong> {safe_html(personas)}<br>
            <strong>Discovery prompt:</strong> {safe_html(first_question)}
        </div>
        """, unsafe_allow_html=True)


def render_card(card, data_source_label, view_mode="Quick Brief"):
    priority_class = card.get("priority", "Medium").lower().replace(" ", "-")
    badges = "".join(badge_html(tag) for tag in card.get("tags", []))
    compact = card.get("compact", {})

    st.markdown(f"""
        <div class="card">
            <h3>{safe_html(card['name'])} <span class="priority priority-{priority_class}">{safe_html(card.get('priority', ''))}</span></h3>
            <div class="meta">
                {safe_html(card.get('tier', ''))}
                | Overall Score {safe_html(card.get('score', ''))}
                | Strategic Score {safe_html(card.get('strategic_score', ''))}
                | Enrollment {safe_html(card.get('enrollment', ''))}
                | Score Band {safe_html(card.get('score_band', ''))}
                | Signal Density {safe_html(card.get('signal_density', ''))}
            </div>
            <div>{badges}</div>
        </div>
    """, unsafe_allow_html=True)

    render_summary_grid(card)
    render_quick_prep(card)
    render_pcg_solution_matches(card)

    st.markdown("**Top Signals**")
    for item in compact.get("top_signals", [])[:3]:
        st.markdown(f"- {item}")

    st.markdown("**Why This District Scored This Way**")
    for item in card.get("score_explanation", [])[:4]:
        st.markdown(f"- {item}")

    st.markdown("**Recommended Next Moves**")
    for item in card.get("next_moves", [])[:4]:
        st.markdown(f"- {item}")

    st.markdown("**Discovery Questions**")
    for item in card.get("questions", [])[:5]:
        st.markdown(f"- {item}")

    contacts = card.get("contacts", [])
    st.markdown("**Key Contacts**")

    if contacts:
        for item in contacts[:4]:
            st.markdown(f"- {item}")

        if len(contacts) > 4:
            with st.expander("Show all contacts", expanded=False):
                for item in contacts[4:8]:
                    st.markdown(f"- {item}")
    else:
        st.markdown("_No contacts were available for this district._")

    if view_mode == "Full Detail":
        with st.expander("Deeper Planning Detail", expanded=False):
            st.markdown("**Score Breakdown**")
            for item in card.get("score_breakdown", []):
                st.markdown(f"- {item}")

            st.markdown("**Strategic Signals**")
            for item in card.get("signals", []):
                st.markdown(f"- {item}")

            st.markdown("**PCG / Emerald Alignment**")
            for item in card.get("alignment", []):
                st.markdown(f"- {item}")

            st.markdown("**Full Listen For**")
            st.markdown("".join(chip_html(item) for item in card.get("listen", [])), unsafe_allow_html=True)

            st.markdown("**Full Guidance**")
            for item in card.get("avoid", []):
                st.markdown(f"- {item}")

    quick_prep = (
        f"{card['name']} Quick Prep\n\n"
        f"Conversation starter: {card.get('conversation_starter', '')}\n\n"
        f"Why it matters: {compact.get('why', '')}\n"
        f"Best entry point: {compact.get('entry', '')}\n"
        f"Likely barrier: {compact.get('barrier', '')}\n\n"
        f"Top PCG solution pathways:\n"
        + "\n".join([f"- {m.get('solution','')} ({m.get('fit_level','')}): {m.get('positioning','')}" for m in card.get("pcg_service_matches", [])[:3]])
        + "\n\n"
        f"Best opening question: {(card.get('questions') or [''])[0]}\n\n"
        f"Listen for: {', '.join(card.get('listen', [])[:8])}\n\n"
        f"Recommended next moves: {'; '.join(card.get('next_moves', [])[:3])}"
    )

    with st.expander("Copy Quick Prep", expanded=False):
        st.text_area("Quick prep copy", quick_prep, height=260, key=f"quick_{card['name']}")

    st.download_button(
        "Download this district profile",
        data=build_docx([card], data_source_label=data_source_label),
        file_name=f"{card['name'].replace(' ', '_')}_District_Profile.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key=f"download_{card['name']}",
    )


def render_hot_accounts(cards):
    st.markdown('<div class="section-title">Hot Accounts</div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="helper-text">Top opportunities based on priority, score, signal density, relationship context, and PCG solution fit.</div>',
        unsafe_allow_html=True,
    )

    if not cards:
        st.info("No hot accounts available.")
        return

    hot = sorted(
        cards,
        key=lambda c: (
            priority_rank(c.get("priority")),
            score_sort_value(c),
            len(c.get("pcg_service_matches", [])),
            1 if "Existing Relationship" in c.get("tags", []) else 0,
            len(c.get("signals", [])),
        ),
        reverse=True,
    )[:8]

    for c in hot:
        top_solution = c.get("pcg_service_matches", [{}])[0].get("solution", "Discovery needed") if c.get("pcg_service_matches") else "Discovery needed"
        tags = ", ".join(c.get("tags", [])[:4])
        st.markdown(f"""
        <div class="hot-card">
            <strong>{safe_html(c.get("name", ""))}</strong> — {safe_html(c.get("priority", ""))}<br>
            <span class="helper-text">Score {safe_html(c.get("score", ""))} | {safe_html(c.get("tier", ""))} | {safe_html(tags)}</span><br>
            <strong>Top pathway:</strong> {safe_html(top_solution)}<br>
            {safe_html(c.get("compact", {}).get("why", ""))}
        </div>
        """, unsafe_allow_html=True)


def render_opportunity_matrix(cards):
    st.markdown('<div class="section-title">Matrix View</div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="helper-text">A short in-app matrix for quick scanning. Use the Word download for the full matrix.</div>',
        unsafe_allow_html=True,
    )

    if not cards:
        st.info("No cards to show.")
        return

    df = pd.DataFrame([
        {
            "District": c.get("name", ""),
            "Tier": c.get("tier", ""),
            "Priority": c.get("priority", ""),
            "Score": c.get("score", ""),
            "Strategic Score": c.get("strategic_score", ""),
            "Score Band": c.get("score_band", ""),
            "Signal Density": c.get("signal_density", ""),
            "Top PCG Pathway": c.get("pcg_service_matches", [{}])[0].get("solution", "") if c.get("pcg_service_matches") else "",
            "Why It Matters": c.get("compact", {}).get("why", ""),
            "Best Entry Point": c.get("compact", {}).get("entry", ""),
            "Conversation Starter": c.get("conversation_starter", ""),
            "Contacts": "; ".join(c.get("contacts", [])[:3]),
        }
        for c in cards
    ])

    st.dataframe(df, use_container_width=True, hide_index=True)

    st.download_button(
        "Download short matrix as CSV",
        data=df.to_csv(index=False).encode("utf-8"),
        file_name="district_opportunity_matrix.csv",
        mime="text/csv",
    )


def render_compare_mode(cards):
    st.markdown('<div class="section-title">Compare Districts</div>', unsafe_allow_html=True)

    if len(cards) < 2:
        st.info("Select at least two visible districts to compare.")
        return

    options = [c["name"] for c in cards]

    selected = st.multiselect(
        "Choose districts to compare",
        options=options,
        default=options[: min(3, len(options))],
        max_selections=4,
    )

    selected_cards = [c for c in cards if c["name"] in selected]

    if not selected_cards:
        return

    compare_df = pd.DataFrame([
        {
            "District": c.get("name", ""),
            "Tier": c.get("tier", ""),
            "Priority": c.get("priority", ""),
            "Score": c.get("score", ""),
            "Strategic Score": c.get("strategic_score", ""),
            "Top Tags": ", ".join(c.get("tags", [])[:5]),
            "Top PCG Pathway": c.get("pcg_service_matches", [{}])[0].get("solution", "") if c.get("pcg_service_matches") else "",
            "Best Entry Point": c.get("compact", {}).get("entry", ""),
            "Likely Barrier": c.get("compact", {}).get("barrier", ""),
            "Next Move": "; ".join(c.get("next_moves", [])[:2]),
        }
        for c in selected_cards
    ])

    st.dataframe(compare_df, use_container_width=True, hide_index=True)

    common_tags = set(selected_cards[0].get("tags", []))
    for c in selected_cards[1:]:
        common_tags = common_tags.intersection(set(c.get("tags", [])))

    if common_tags:
        st.markdown("**Common strategic themes:** " + ", ".join(sorted(common_tags)))
    else:
        st.markdown("**Common strategic themes:** No exact tag overlap; compare individual entry points.")


def render_how_to(data_source_label):
    st.markdown('<div class="section-title">How-To Guide</div>', unsafe_allow_html=True)

    st.markdown(f"""
    <div class="howto-box">
    <strong>Current data source:</strong> {safe_html(data_source_label)}<br>
    The app opens with the built-in workbook from GitHub when available. Uploading a workbook in the sidebar temporarily overrides the built-in data for that session.
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
    <div class="howto-box">
    <strong>60-second BD workflow</strong>
    <ol>
      <li>Search the district name.</li>
      <li>Read Why It Matters, Best Entry Point, and Likely Barrier.</li>
      <li>Review Recommended PCG Solution Pathways.</li>
      <li>Use the Quick Prep conversation starter to open discovery.</li>
      <li>Review decision-makers and influencers.</li>
      <li>Download the District Profile Document for team follow-up.</li>
    </ol>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("### Useful searches")
    st.code("Dallas\nFort Worth\nSPED\nMTSS\nEureka\nBluebonnet\ncareer\nteacher burden\nliteracy\ndyslexia\ncoaching\nPlaybook", language="text")

    st.markdown("### Tag legend")

    legend = {
        "Math": "Math growth, numeracy, STAAR math, Algebra readiness, student practice.",
        "MTSS": "Intervention, Tier 2/Tier 3, progress monitoring, campus variation.",
        "SPED/ELL": "Special education, multilingual learners, subgroup access, differentiated supports.",
        "Literacy / Dyslexia": "Science of Reading, structured literacy, dyslexia, reading intervention, literacy audit/review.",
        "Coaching / Leadership": "Instructional coaching, leadership development, walkthroughs, observation/feedback, coaching infrastructure.",
        "CCMR": "College, career, and military readiness; CTE; pathways; career awareness.",
        "Teacher Capacity": "Professional learning, coaching, PLCs, teacher burden, implementation support.",
        "Curriculum / HQIM": "Curriculum adoption, HQIM, Bluebonnet, Eureka, implementation fidelity.",
        "Data Systems": "Dashboards, analytics, progress monitoring, implementation monitoring.",
        "Funding / Grants": "Title funds, grants, Medicaid, funding alignment, sustainability.",
        "Existing Relationship": "A relationship or contract signal exists in the workbook.",
    }

    for tag, meaning in legend.items():
        st.markdown(f"{badge_html(tag)} {meaning}", unsafe_allow_html=True)


def show_workbook_debug(
    sheet_names,
    strategic_df,
    scorecard_df,
    basic_df,
    contacts_df,
    leadership_df,
    csi_df,
    tsi_df,
    audit_df,
    score_cols,
    data_source_label,
):
    st.markdown('<div class="section-title">Workbook Diagnostics</div>', unsafe_allow_html=True)

    st.markdown(
        f'<div class="helper-text">Current data source: <strong>{safe_html(data_source_label)}</strong></div>',
        unsafe_allow_html=True,
    )

    st.markdown("### Sheets found")
    st.write(sheet_names)

    st.markdown("### Detected scorecard columns")
    st.write(score_cols)

    for label, df in [
        ("Strategic Indicators", strategic_df),
        ("Master Scorecard", scorecard_df),
        ("Basic District Info", basic_df),
        ("District Contacts", contacts_df),
        ("Leadership and Governance", leadership_df),
        ("CSI", csi_df),
        ("TSI", tsi_df),
    ]:
        with st.expander(label, expanded=False):
            st.write(f"Rows: {len(df)}")
            st.write("Columns:")
            st.write(list(df.columns))

            if DISTRICT_COL in df.columns:
                st.write("First district names found:")
                st.write(df[DISTRICT_COL].dropna().astype(str).head(15).tolist())

    st.markdown("### Eligibility audit")
    st.dataframe(audit_df, use_container_width=True)


# ============================================================
# APP LAYOUT
# ============================================================

st.markdown('<div class="main-title">Strategic District Field Guide</div>', unsafe_allow_html=True)
st.markdown(
    '<div class="subtitle">BD-ready district profiles with PCG service pathway recommendations, Emerald alignment, and leadership approach guidance.</div>',
    unsafe_allow_html=True,
)

if "shortlist" not in st.session_state:
    st.session_state.shortlist = []

with st.sidebar:
    st.header("1. Data")
    uploaded_file = st.file_uploader("Optional: upload workbook override", type=["xlsx"])
    st.caption("If no workbook is uploaded, the app will use the built-in workbook from GitHub.")

workbook_payload, data_source_label, source_type = get_workbook_source(uploaded_file)

if workbook_payload is None:
    tab_field, tab_hot, tab_compare, tab_howto, tab_matrix, tab_diag = st.tabs(
        ["Field Guide", "Hot Accounts", "Compare Districts", "How-To Guide", "Matrix View", "Workbook Diagnostics"]
    )

    with tab_field:
        st.error(f"No workbook found. Add {BUILT_IN_WORKBOOK} to the GitHub repo or upload a workbook in the sidebar.")

    with tab_howto:
        render_how_to(data_source_label)

    with tab_matrix:
        st.info("Workbook required to populate the matrix.")

    with tab_diag:
        st.info("Workbook required to view diagnostics.")

    st.stop()

with st.spinner("Loading district intelligence..."):
    (
        cards,
        sheet_names,
        strategic_df,
        scorecard_df,
        basic_df,
        contacts_df,
        leadership_df,
        csi_df,
        tsi_df,
        audit_df,
        score_cols,
    ) = load_cards_from_workbook_cached(workbook_payload, source_type)

if not cards:
    tab_field, tab_hot, tab_compare, tab_howto, tab_matrix, tab_diag = st.tabs(
        ["Field Guide", "Hot Accounts", "Compare Districts", "How-To Guide", "Matrix View", "Workbook Diagnostics"]
    )

    with tab_field:
        st.warning("No eligible district cards were generated. Review the Workbook Diagnostics tab.")

    with tab_howto:
        render_how_to(data_source_label)

    with tab_matrix:
        st.info("No eligible cards available.")

    with tab_diag:
        show_workbook_debug(
            sheet_names,
            strategic_df,
            scorecard_df,
            basic_df,
            contacts_df,
            leadership_df,
            csi_df,
            tsi_df,
            audit_df,
            score_cols,
            data_source_label,
        )

    st.stop()

all_districts = [card.get("name", "") for card in cards]
all_tiers = sorted({card.get("tier", "") for card in cards if card.get("tier")})
all_tags = sorted({tag for card in cards for tag in card.get("tags", [])})

with st.sidebar:
    st.success(f"Using: {data_source_label}")

    st.header("2. Search")
    query = st.text_input("Search", placeholder="District, contact, signal, offering...")

    st.header("3. Filters")
    selected_tiers = st.multiselect("Tier", options=all_tiers, default=[])
    selected_tags = st.multiselect("Strategic / solution tags", options=all_tags, default=[])

    st.header("4. Shortlist")
    shortlist_input = st.multiselect(
        "Shortlist districts",
        options=all_districts,
        default=[x for x in st.session_state.shortlist if x in all_districts],
    )
    st.session_state.shortlist = shortlist_input

    show_shortlist_only = st.checkbox("Show shortlisted only", value=False)

    if st.button("Clear shortlist"):
        st.session_state.shortlist = []
        st.rerun()

    st.header("5. Display")
    view_mode = st.radio("View mode", options=["Quick Brief", "Full Detail"], index=0)

    sort_mode = st.selectbox(
        "Sort by",
        options=[
            "Priority then Score",
            "Overall Score",
            "Strategic Score",
            "Enrollment",
            "Signal Density",
            "District A-Z",
        ],
        index=0,
    )

    st.caption("Tip: Use Quick Brief for live conversations. Download the District Profile Document for team follow-up.")

filtered_cards = filter_cards(
    cards,
    query,
    selected_tiers,
    selected_tags,
    st.session_state.shortlist,
    show_shortlist_only,
)

filtered_cards = sort_cards(filtered_cards, sort_mode)

tab_field, tab_hot, tab_compare, tab_howto, tab_matrix, tab_diag = st.tabs(
    ["Field Guide", "Hot Accounts", "Compare Districts", "How-To Guide", "Matrix View", "Workbook Diagnostics"]
)

with tab_field:
    metric1, metric2, metric3, metric4, metric5 = st.columns(5)

    metric1.metric("📍 Cards Shown", len(filtered_cards))
    metric2.metric("✅ Eligible", len(cards))
    metric3.metric(
        "🔥 High Priority",
        sum(1 for card in filtered_cards if card.get("priority") in ["High", "Very High"]),
    )
    metric4.metric(
        "⭐ Tier 1",
        sum(1 for card in filtered_cards if card.get("tier") == "Tier 1"),
    )
    metric5.metric(
        "🧠 High Signal",
        sum(1 for card in filtered_cards if card.get("signal_density") == "High"),
    )

    if query:
        st.markdown(
            f'<div class="helper-text">Showing results for: <strong>{safe_html(query)}</strong></div>',
            unsafe_allow_html=True,
        )

    col_dl1, col_dl2, col_dl3 = st.columns([1, 1, 1])

    with col_dl1:
        st.download_button(
            "Download District Profile Document",
            data=build_docx(filtered_cards, data_source_label=data_source_label),
            file_name="District_Profile_Document.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            disabled=not filtered_cards,
        )

    with col_dl2:
        st.download_button(
            "Download full matrix",
            data=build_matrix_docx(filtered_cards, data_source_label=data_source_label),
            file_name="Texas_District_Strategic_Positioning_Matrix.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            disabled=not filtered_cards,
        )

    with col_dl3:
        selected_brief = st.selectbox(
            "One-district profile",
            options=[""] + [c["name"] for c in filtered_cards],
        )

        if selected_brief:
            selected_card = next(c for c in filtered_cards if c["name"] == selected_brief)

            st.download_button(
                "Download selected district profile",
                data=build_docx([selected_card], data_source_label=data_source_label),
                file_name=f"{selected_brief.replace(' ', '_')}_District_Profile.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download_selected_brief",
            )

    st.divider()

    if not filtered_cards:
        st.info("No district cards match the current search/filter.")
    else:
        for card in filtered_cards:
            left, right = st.columns([5, 1])

            with left:
                render_card(card, data_source_label=data_source_label, view_mode=view_mode)

            with right:
                district_name = card.get("name")

                if district_name in st.session_state.shortlist:
                    if st.button("Remove ⭐", key=f"remove_{district_name}"):
                        st.session_state.shortlist = [
                            x for x in st.session_state.shortlist if x != district_name
                        ]
                        st.rerun()
                else:
                    if st.button("Shortlist ⭐", key=f"add_{district_name}"):
                        st.session_state.shortlist.append(district_name)
                        st.rerun()

            st.divider()

with tab_hot:
    render_hot_accounts(filtered_cards)

with tab_compare:
    render_compare_mode(filtered_cards)

with tab_howto:
    render_how_to(data_source_label)

with tab_matrix:
    render_opportunity_matrix(filtered_cards)

    st.download_button(
        "Download full matrix as Word document",
        data=build_matrix_docx(filtered_cards, data_source_label=data_source_label),
        file_name="Texas_District_Strategic_Positioning_Matrix.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        disabled=not filtered_cards,
        key="matrix_tab_download",
    )

with tab_diag:
    show_workbook_debug(
        sheet_names,
        strategic_df,
        scorecard_df,
        basic_df,
        contacts_df,
        leadership_df,
        csi_df,
        tsi_df,
        audit_df,
        score_cols,
        data_source_label,
    )
